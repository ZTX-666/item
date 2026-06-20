# -*- coding: utf-8 -*-
"""
Build Site Memorandum page 1 (.docx) with mandatory company letterhead in the header.

Letterhead image path (default): <skill>/assets/letterhead.png
Run from anywhere: python generate_site_memo_docx.py --help
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

_SCRIPTS = Path(__file__).resolve().parent
if str(_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS))

from contact_db import ContactDB, default_db_path, resolve_media_path

from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def skill_root() -> Path:
    return Path(__file__).resolve().parents[1]


def default_letterhead() -> Path:
    return skill_root() / "assets" / "letterhead.png"


def add_letterhead_header(doc: Document, letterhead: Path | None, width_inches: float = 6.2) -> None:
    """Insert PNG into primary section header. If file missing, add visible reminder in header."""
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    if letterhead and letterhead.is_file():
        run = hp.add_run()
        run.add_picture(str(letterhead), width=Inches(width_inches))
    else:
        hp.add_run(
            "【页眉：请将公司 logo 存为本技能目录 assets/letterhead.png 后重新生成】"
        ).font.size = Pt(9)


@dataclass(frozen=True)
class MemoRoutingBlock:
    """Header grid matching the standard Word template (4 cols × 4 rows + 2 merged rows)."""

    our_ref: str
    date: str
    your_ref: str
    pages: str
    to: str
    attn: str
    copy_to: str
    fax: str
    project: str
    re: str


def _set_cell_text_font(cell, text: str, *, bold: bool = False, size_pt: int = 11) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for r in paragraph.runs:
            r.bold = bold
            r.font.size = Pt(size_pt)


def add_standard_four_column_routing_table(doc: Document, v: MemoRoutingBlock) -> None:
    """
    Standard Site Memorandum routing block:
    - Rows 1–4: four columns [label | value | label | value].
    - Rows 5–6: label in col 1; cols 2–4 merged for long text (Project, Re).
    Full grid borders via built-in style when available.
    """
    table = doc.add_table(rows=6, cols=4)
    table.autofit = False
    for name in ("Table Grid", "网格型", "网格式"):
        try:
            table.style = name
            break
        except KeyError:
            continue

    # Row 1: Our Ref | value | Date | value
    r0 = table.rows[0].cells
    _set_cell_text_font(r0[0], "Our Ref. :", bold=True)
    _set_cell_text_font(r0[1], v.our_ref)
    _set_cell_text_font(r0[2], "Date :", bold=True)
    _set_cell_text_font(r0[3], v.date)

    # Row 2: Your Ref | value | Pages | value
    r1 = table.rows[1].cells
    _set_cell_text_font(r1[0], "Your Ref. :", bold=True)
    _set_cell_text_font(r1[1], v.your_ref)
    _set_cell_text_font(r1[2], "Pages :", bold=True)
    _set_cell_text_font(r1[3], v.pages)

    # Row 3: To | value | Attn. | value
    r2 = table.rows[2].cells
    _set_cell_text_font(r2[0], "To :", bold=True)
    _set_cell_text_font(r2[1], v.to)
    _set_cell_text_font(r2[2], "Attn. :", bold=True)
    _set_cell_text_font(r2[3], v.attn)

    # Row 4: Copy To | value | Fax No. | value
    r3 = table.rows[3].cells
    _set_cell_text_font(r3[0], "Copy To :", bold=True)
    _set_cell_text_font(r3[1], v.copy_to)
    _set_cell_text_font(r3[2], "Fax No. :", bold=True)
    _set_cell_text_font(r3[3], v.fax)

    # Row 5: Project — merge cols 2–4 (index 1..3)
    r4 = table.rows[4].cells
    _set_cell_text_font(r4[0], "Project :", bold=True)
    r4[1].merge(r4[3])
    _set_cell_text_font(r4[1], v.project)

    # Row 6: Re — merge cols 2–4
    r5 = table.rows[5].cells
    _set_cell_text_font(r5[0], "Re :", bold=True)
    r5[1].merge(r5[3])
    _set_cell_text_font(r5[1], v.re)

    # Column widths (approx. A4 printable ~6.5 in for body table)
    w_label_narrow = Inches(0.95)
    w_value_wide = Inches(2.05)
    for row in table.rows[:4]:
        row.cells[0].width = w_label_narrow
        row.cells[1].width = w_value_wide
        row.cells[2].width = w_label_narrow
        row.cells[3].width = w_value_wide
    for row in table.rows[4:]:
        row.cells[0].width = w_label_narrow
        row.cells[1].width = Inches(2.05 + 0.95 + 2.05)


def build_document(
    *,
    letterhead: Path | None,
    routing: MemoRoutingBlock,
    body: str,
    responsibility: str,
    company_line: str,
    signatory_name: str,
    signatory_title: str,
    signature_image_path: Path | None = None,
    signature_width_inches: float = 2.2,
) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    # Header must contain letterhead (mandatory for this skill)
    add_letterhead_header(doc, letterhead)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SITE MEMORANDUM")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph()

    add_standard_four_column_routing_table(doc, routing)

    doc.add_paragraph()
    bp = doc.add_paragraph()
    bp.paragraph_format.space_after = Pt(6)
    bp.add_run(body).font.size = Pt(11)

    # Footer block (match standard memo / PDF): 责任 → 公司名 → 电子签名图 → 横线 → 姓名 → 职位
    resp = doc.add_paragraph()
    resp.paragraph_format.space_after = Pt(3)
    resp.add_run("责任: ").bold = True
    resp.add_run(responsibility).font.size = Pt(11)

    comp = doc.add_paragraph()
    comp.paragraph_format.space_before = Pt(2)
    comp.paragraph_format.space_after = Pt(2)
    comp.add_run(company_line).font.size = Pt(11)

    if signature_image_path is not None and signature_image_path.is_file():
        sig_img = doc.add_paragraph()
        sig_img.paragraph_format.space_before = Pt(0)
        sig_img.paragraph_format.space_after = Pt(2)
        run_img = sig_img.add_run()
        run_img.add_picture(str(signature_image_path), width=Inches(signature_width_inches))

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(3)
    line.add_run("_" * 42).font.size = Pt(11)

    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(2)
    sig.paragraph_format.space_after = Pt(0)
    sig.add_run(signatory_name).font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title.add_run(signatory_title).font.size = Pt(11)
    return doc


def _resolve_contacts(db_path: Path | None, field: str | None, literal: str, style: str) -> str:
    if not field or not str(field).strip():
        return literal
    db = ContactDB(db_path).load()
    return db.resolve_tokens([field], style=style)


def main() -> None:
    ap = argparse.ArgumentParser(description="Site Memorandum page 1 (.docx) with header letterhead")
    ap.add_argument("--output", "-o", type=Path, default=Path.cwd() / "SITE-MEMORANDUM.docx")
    ap.add_argument("--letterhead", type=Path, default=None, help="PNG path (default: skill assets/letterhead.png)")
    ap.add_argument(
        "--contacts-db",
        type=Path,
        default=None,
        help="Override path to contacts.xlsx (default: skill data/contacts.xlsx; legacy .json supported)",
    )
    ap.add_argument(
        "--contact-style",
        choices=("name_title", "name_only"),
        default="name_title",
        help="How resolved contacts render in To/Copy To/Attn.",
    )
    ap.add_argument(
        "--to-contact",
        default=None,
        help="Abbrev(s): comma/space-separated; overrides --to when set",
    )
    ap.add_argument("--copy-to-contact", default=None, help="Abbrev(s); overrides --copy-to when set")
    ap.add_argument("--attn-contact", default=None, help="Abbrev(s); overrides --attn when set")
    ap.add_argument("--our-ref", default="EI-2026-0001")
    ap.add_argument("--date", default="19 Apr 2026")
    ap.add_argument("--your-ref", default="N/A")
    ap.add_argument("--pages", default="Page 1 of 1")
    ap.add_argument("--to", dest="to_", default="Xu Siyuan")
    ap.add_argument("--attn", default="-")
    ap.add_argument("--copy-to", default="He Wenkai")
    ap.add_argument("--fax", default="N/A")
    ap.add_argument("--project", default="FEW")
    ap.add_argument("--re", dest="re_", default="Site Memo Test")
    ap.add_argument("--body", default="Site Memo Test.")
    ap.add_argument(
        "--body-file",
        type=Path,
        default=None,
        help="UTF-8 file for body (field 11); when set, overrides --body (avoids shell quoting)",
    )
    ap.add_argument("--responsibility", default="（待工程师确认）")
    ap.add_argument(
        "--company",
        default="",
        help="Footer company line; if empty and --signatory-contact is set, uses that contact's company column",
    )
    ap.add_argument("--signatory", default="Xu Siyuan")
    ap.add_argument("--title", default="ASM")
    ap.add_argument(
        "--signatory-contact",
        default=None,
        help="Contact abbrev: sets printed name/title from DB and uses that row's signature_image",
    )
    ap.add_argument(
        "--signature-image",
        default=None,
        help="E-signature image path (absolute or relative to skill root). Ignored if --signatory-contact is set.",
    )
    ap.add_argument(
        "--signature-width-inches",
        type=float,
        default=2.2,
        help="Rendered width of e-signature image in the body",
    )
    args = ap.parse_args()

    body_text = (
        args.body_file.read_text(encoding="utf-8").strip()
        if args.body_file is not None
        else args.body
    )

    lh = args.letterhead if args.letterhead else default_letterhead()
    dbp = args.contacts_db if args.contacts_db else default_db_path()
    skill = skill_root()
    try:
        to_val = _resolve_contacts(dbp, args.to_contact, args.to_, args.contact_style)
        attn_val = _resolve_contacts(dbp, args.attn_contact, args.attn, args.contact_style)
        copy_val = _resolve_contacts(dbp, args.copy_to_contact, args.copy_to, args.contact_style)
    except LookupError as e:
        raise SystemExit(f"Contact resolve error: {e}") from e

    sign_name = args.signatory
    sign_title = args.title
    sig_path: Path | None = None
    company_line = (args.company or "").strip()
    if args.signatory_contact:
        db_sig = ContactDB(dbp).load()
        c_sig = db_sig.find(str(args.signatory_contact).strip())
        if not c_sig:
            raise SystemExit(f"Unknown signatory-contact: {args.signatory_contact!r}")
        sign_name = c_sig.memo_english_name()
        sign_title = c_sig.title
        sig_path = resolve_media_path(skill, c_sig.signature_image)
        if not company_line:
            company_line = (c_sig.company or "").strip() or "（公司名称待补）"
    elif args.signature_image:
        sig_path = resolve_media_path(skill, str(args.signature_image).strip())
    if not args.signatory_contact and not company_line:
        company_line = "（公司名称待补）"

    routing = MemoRoutingBlock(
        our_ref=args.our_ref,
        date=args.date,
        your_ref=args.your_ref,
        pages=args.pages,
        to=to_val,
        attn=attn_val,
        copy_to=copy_val,
        fax=args.fax,
        project=args.project,
        re=args.re_,
    )
    doc = build_document(
        letterhead=lh,
        routing=routing,
        body=body_text,
        responsibility=args.responsibility,
        company_line=company_line,
        signatory_name=sign_name,
        signatory_title=sign_title,
        signature_image_path=sig_path,
        signature_width_inches=args.signature_width_inches,
    )
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    print(str(args.output.resolve()))


if __name__ == "__main__":
    main()
