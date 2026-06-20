import json
import sys
from pathlib import Path


def replace_paragraph_text_preserve_runs(paragraph, new_text: str) -> None:
    """
    Best-effort text replacement while preserving formatting:
    - If paragraph has exactly one run, replace that run's text.
    - Else, fall back to paragraph.text (may normalize runs).
    """
    runs = list(paragraph.runs)
    if len(runs) == 1:
        runs[0].text = new_text
        return
    paragraph.text = new_text


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def main() -> int:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        print("ERROR: python-docx is not installed. Install with: python -m pip install python-docx", file=sys.stderr)
        print(str(e), file=sys.stderr)
        return 2

    if len(sys.argv) < 4:
        print("Usage: python apply_translation.py <input.docx> <mapping.json> <output.docx>", file=sys.stderr)
        return 2

    src = Path(sys.argv[1]).expanduser().resolve()
    mapping_path = Path(sys.argv[2]).expanduser().resolve()
    out = Path(sys.argv[3]).expanduser().resolve()

    mapping = json.loads(mapping_path.read_text(encoding="utf-8"))
    translations: dict[str, str] = mapping.get("translations", {})

    doc = Document(str(src))

    replaced = 0
    for p in iter_all_paragraphs(doc):
        old = p.text or ""
        if old in translations:
            replace_paragraph_text_preserve_runs(p, translations[old])
            replaced += 1

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Replaced {replaced} paragraphs/cell-paragraphs. Saved: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

