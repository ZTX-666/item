import json
import sys
from pathlib import Path


def iter_paragraphs_in_cell(cell):
    # python-docx cell.paragraphs already covers paragraphs in the cell.
    for p in cell.paragraphs:
        yield p


def main() -> int:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        print("ERROR: python-docx is not installed. Install with: python -m pip install python-docx", file=sys.stderr)
        print(str(e), file=sys.stderr)
        return 2

    if len(sys.argv) < 3:
        print("Usage: python extract_docx_text.py <input.docx> <output.json>", file=sys.stderr)
        return 2

    src = Path(sys.argv[1]).expanduser().resolve()
    out = Path(sys.argv[2]).expanduser().resolve()

    if not src.exists():
        print(f"ERROR: input file not found: {src}", file=sys.stderr)
        return 2

    doc = Document(str(src))

    items = []
    item_id = 0

    # Body paragraphs
    for p in doc.paragraphs:
        text = p.text or ""
        if text.strip() == "":
            continue
        items.append(
            {
                "id": f"p{item_id}",
                "kind": "paragraph",
                "text": text,
            }
        )
        item_id += 1

    # Tables
    t_idx = 0
    for t in doc.tables:
        r_idx = 0
        for row in t.rows:
            c_idx = 0
            for cell in row.cells:
                # cell.text merges all paragraphs; we'll store by paragraph too to keep mapping stable.
                for p_idx, p in enumerate(iter_paragraphs_in_cell(cell)):
                    text = p.text or ""
                    if text.strip() == "":
                        continue
                    items.append(
                        {
                            "id": f"t{t_idx}r{r_idx}c{c_idx}p{p_idx}",
                            "kind": "table_cell_paragraph",
                            "text": text,
                        }
                    )
                    item_id += 1
                c_idx += 1
            r_idx += 1
        t_idx += 1

    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(json.dumps({"source": str(src), "items": items}, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Wrote {len(items)} items to {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

