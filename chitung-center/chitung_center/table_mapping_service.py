from __future__ import annotations

import json
import re
import subprocess
import tempfile
import time
import uuid
from pathlib import Path
from typing import Any

from chitung_center.config import settings


class TableMappingError(RuntimeError):
    pass


def list_table_mapping_forms() -> dict[str, Any]:
    script_dir = _script_dir()
    forms = _load_forms(script_dir)
    return {
        "ok": True,
        "script_dir": str(script_dir),
        "script_available": _script_available(script_dir),
        "items": sorted(forms.values(), key=lambda item: _form_sort_key(str(item.get("id", "")))),
    }


def get_table_mapping_form(form_id: str) -> dict[str, Any]:
    forms = _load_forms(_script_dir())
    form = forms.get(form_id)
    if not form:
        raise TableMappingError(f"Unknown table mapping form: {form_id}")
    return {"ok": True, "item": form}


def extract_table_mapping_fields(file_path: str, form_id: str) -> dict[str, Any]:
    script_dir = _script_dir()
    form = get_table_mapping_form(form_id)["item"]
    document = _read_docx(Path(file_path))
    fields = _extract_fields(document, form)
    return {
        "ok": True,
        "form": form,
        "source_path": str(Path(file_path)),
        "fields": fields,
        "matched_count": sum(1 for item in fields.values() if item.get("value")),
        "document_preview": document["preview"],
        "script_available": _script_available(script_dir),
    }


def run_table_mapping_fill(
    *,
    file_path: str,
    form_id: str,
    fields: dict[str, str] | None = None,
    action: str = "draft",
    screenshot: bool = True,
    dry_run: bool = False,
) -> dict[str, Any]:
    script_dir = _script_dir()
    form = get_table_mapping_form(form_id)["item"]
    extracted = extract_table_mapping_fields(file_path, form_id)
    payload = _clean_field_payload(fields or {key: item.get("value", "") for key, item in extracted["fields"].items()})

    command = [settings.table_mapping_node_bin, "fill.js", form_id]
    if action not in {"draft", "save_draft"}:
        raise TableMappingError("Only draft/save_draft is supported in the first table mapping integration.")

    if not dry_run and not _script_available(script_dir):
        raise TableMappingError(f"Table mapping script is not available: {script_dir}")

    job_id = uuid.uuid4().hex
    command_preview = command + ["<fields.json>", "--draft"]
    if screenshot:
        command_preview.append("--screenshot")

    result: dict[str, Any] = {
        "ok": True,
        "job_id": job_id,
        "dry_run": dry_run,
        "form": form,
        "source_path": str(Path(file_path)),
        "fields": payload,
        "extracted": extracted["fields"],
        "command": command_preview,
    }
    if dry_run:
        return result

    with tempfile.TemporaryDirectory(prefix="chitung_table_mapping_") as tmp_dir:
        data_path = Path(tmp_dir) / "fields.json"
        data_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        run_command = command + [str(data_path), "--draft"]
        if screenshot:
            run_command.append("--screenshot")

        started = time.time()
        completed = subprocess.run(
            run_command,
            cwd=str(script_dir),
            text=True,
            encoding="utf-8",
            errors="replace",
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=max(30, int(settings.table_mapping_timeout_seconds)),
            shell=False,
        )
        duration_ms = int((time.time() - started) * 1000)

    result.update(
        {
            "ok": completed.returncode == 0,
            "exit_code": completed.returncode,
            "duration_ms": duration_ms,
            "stdout": completed.stdout,
            "stderr": completed.stderr,
        }
    )
    return result


def _script_dir() -> Path:
    return Path(settings.table_mapping_script_dir)


def _script_available(script_dir: Path) -> bool:
    return (script_dir / "fill.js").exists() and (script_dir / "login.js").exists()


def _load_forms(script_dir: Path) -> dict[str, dict[str, Any]]:
    forms: dict[str, dict[str, Any]] = {}
    config_path = script_dir / "config.json"
    if config_path.exists():
        data = json.loads(config_path.read_text(encoding="utf-8"))
        for form_id, item in data.get("forms", {}).items():
            fields = item.get("fields", {}) if isinstance(item, dict) else {}
            forms[str(form_id)] = {
                "id": str(form_id),
                "name": item.get("name", str(form_id)),
                "category": item.get("category", ""),
                "fields": [{"name": name, "type": "text", "required": False} for name in fields.keys()],
            }

    explore_path = script_dir.parent / "explore-results-merged.json"
    if explore_path.exists():
        data = json.loads(explore_path.read_text(encoding="utf-8"))
        for form_id, item in data.items():
            field_names = _field_names_from_explore(item)
            if form_id in forms:
                known = {field["name"] for field in forms[form_id]["fields"]}
                forms[form_id]["fields"].extend(
                    {"name": name, "type": "text", "required": False}
                    for name in field_names
                    if name and name not in known
                )
            else:
                forms[str(form_id)] = {
                    "id": str(form_id),
                    "name": item.get("name", str(form_id)),
                    "category": item.get("category", ""),
                    "fields": [{"name": name, "type": "text", "required": False} for name in field_names],
                }
    return forms


def _field_names_from_explore(item: dict[str, Any]) -> list[str]:
    names: list[str] = []
    for field in item.get("fields", []):
        for key in ("label", "ph", "name", "id"):
            value = str(field.get(key, "")).strip()
            cleaned = _clean_field_name(value)
            if cleaned and cleaned not in names:
                names.append(cleaned)
    return names


def _read_docx(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise TableMappingError(f"DOCX file not found: {path}")
    if path.suffix.lower() not in {".docx", ".doc"}:
        raise TableMappingError("Table mapping currently expects a Word document path (.docx preferred).")
    try:
        from docx import Document
    except ImportError as exc:
        raise TableMappingError("python-docx is required to read Word documents.") from exc
    if path.suffix.lower() == ".doc":
        raise TableMappingError("Legacy .doc files are not supported yet. Please save as .docx first.")

    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    preview = "\n".join(paragraphs[:12])
    return {"paragraphs": paragraphs, "tables": tables, "text": "\n".join(paragraphs), "preview": preview}


def _extract_fields(document: dict[str, Any], form: dict[str, Any]) -> dict[str, dict[str, Any]]:
    text = document["text"]
    table_values = _extract_from_tables(document["tables"])
    result: dict[str, dict[str, Any]] = {}
    for field in form.get("fields", []):
        name = str(field.get("name", "")).strip()
        if not name:
            continue
        aliases = _field_aliases(name)
        value = ""
        source = ""
        for alias in aliases:
            if alias in table_values:
                value = table_values[alias]
                source = "table"
                break
        if not value:
            for alias in aliases:
                matched = _match_inline_value(text, alias)
                if matched:
                    value = matched
                    source = "paragraph"
                    break
        result[name] = {"value": value, "source": source, "confidence": 0.85 if value else 0.0}
    return result


def _extract_from_tables(tables: list[list[list[str]]]) -> dict[str, str]:
    values: dict[str, str] = {}
    for rows in tables:
        for row in rows:
            cleaned = [cell.strip() for cell in row if cell and cell.strip()]
            if len(cleaned) < 2:
                continue
            for index, cell in enumerate(cleaned[:-1]):
                key = _clean_field_name(cell)
                value = cleaned[index + 1].strip()
                if key and value and key not in values:
                    values[key] = value
    return values


def _match_inline_value(text: str, field: str) -> str:
    escaped = re.escape(field)
    match = re.search(rf"{escaped}\s*[:：]\s*([^\n\r;；,，]+)", text)
    return match.group(1).strip() if match else ""


def _field_aliases(name: str) -> list[str]:
    cleaned = _clean_field_name(name)
    aliases = [cleaned]
    pairs = [("地盤", "地盘"), ("填報", "填报"), ("檢查", "检查"), ("巡查", "巡查")]
    for source, target in pairs:
        if source in cleaned:
            aliases.append(cleaned.replace(source, target))
        if target in cleaned:
            aliases.append(cleaned.replace(target, source))
    if "日期" in cleaned:
        aliases.append(cleaned.replace("日期", "时间"))
    return list(dict.fromkeys(alias for alias in aliases if alias))


def _clean_field_name(value: str) -> str:
    value = re.sub(r"^(請輸入|请输入|请选择|請選擇)", "", value.strip())
    value = re.sub(r"[:：*＊\s]+$", "", value)
    return value.strip()


def _clean_field_payload(fields: dict[str, str]) -> dict[str, str]:
    return {str(key): str(value).strip() for key, value in fields.items() if str(value).strip()}


def _form_sort_key(form_id: str) -> tuple[int, int, str]:
    match = re.match(r"^(\d+)\.(\d+)$", form_id)
    if not match:
        return (999, 999, form_id)
    return (int(match.group(1)), int(match.group(2)), form_id)
