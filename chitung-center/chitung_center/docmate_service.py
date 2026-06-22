"""DocMate service layer for DOCX upload, changeset preview, commit, and retry."""

from __future__ import annotations

import json
import re
import shutil
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from fastapi import UploadFile

from chitung_center.config import settings
from chitung_center.llm_gateway import llm_gateway
from chitung_center.toolbox_client import toolbox_client


_DOC_CACHE: dict[str, dict[str, Any]] = {}
_CHANGESET_CACHE: dict[str, dict[str, Any]] = {}
_DOWNLOAD_FILES: dict[str, Path] = {}
_VALID_CHANGE_TYPES = {
    "text_replace",
    "text_insert",
    "text_delete",
    "table_cell_update",
    "table_insert",
    "image_insert",
    "image_replace",
}
_VALID_RISK_LEVELS = {"low", "medium", "high"}


async def upload_docx(upload: UploadFile) -> dict[str, Any]:
    """Persist an uploaded DOCX and return a file id for later read/download."""
    filename = _safe_filename(upload.filename or "document.docx")
    suffix = Path(filename).suffix.lower()
    if suffix != ".docx":
        return {"ok": False, "error": "unsupported_format", "summary": "DocMate 仅支持 .docx 文件。"}

    content = await upload.read()
    if not content:
        return {"ok": False, "error": "empty_upload", "summary": "上传文件为空。"}

    file_id = uuid.uuid4().hex
    target = _docmate_upload_dir() / f"{file_id}_{filename}"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_bytes(content)
    _register_download_file(target, file_id=file_id)
    return {
        "ok": True,
        "file_id": file_id,
        "file_name": filename,
        "file_path": str(target),
        "download_url": f"/api/docmate/download/{file_id}",
    }


async def read_docx(file_path: str) -> dict[str, Any]:
    """Step 1: Parse a .docx file into structured paragraphs/tables/images."""
    result = await toolbox_client.call_tool("docmate_read_docx", {"file_path": file_path})
    _remember_doc(file_path, result)
    return result


async def generate_changeset(doc_id: str, instruction: str, context: Any | None = None) -> dict[str, Any]:
    """Step 2: Generate a changeset from natural language, preferring LLM output."""
    llm_result = await _generate_changeset_with_llm(doc_id, instruction, context)
    if llm_result.get("ok"):
        return llm_result

    payload: dict[str, Any] = {"doc_id": doc_id, "instruction": instruction}
    if context is not None:
        payload["context"] = context
    result = await toolbox_client.call_tool("docmate_generate_changeset", payload)
    _remember_toolbox_changeset(result, doc_id, instruction, context)
    if llm_result.get("error"):
        result.setdefault("llm_first_error", llm_result.get("error"))
    return result


async def preview_changeset(changeset_id: str) -> dict[str, Any]:
    """Step 3: Preview the changeset as change cards."""
    cached = _CHANGESET_CACHE.get(changeset_id)
    if cached:
        cards = cached.get("preview_cards") if isinstance(cached.get("preview_cards"), list) else []
        return {
            "ok": True,
            "tool": "docmate_preview_changeset",
            "summary": f"变更预览：共 {len(cards)} 项。",
            "data": {
                "changeset_id": changeset_id,
                "instruction": cached.get("instruction", ""),
                "preview_cards": cards,
                "total_changes": len(cards),
                "generator": cached.get("generator", "llm"),
            },
        }
    return await toolbox_client.call_tool("docmate_preview_changeset", {"changeset_id": changeset_id})


async def apply_changeset(
    changeset_id: str,
    accepted_change_ids: list[str],
    save_as: str | None = None,
) -> dict[str, Any]:
    """Step 4: Apply accepted changes and write the output .docx."""
    cached = _CHANGESET_CACHE.get(changeset_id)
    if cached and cached.get("generator") == "llm" and cached.get("source_path"):
        source_path = Path(str(cached.get("source_path")))
        if source_path.exists():
            return _apply_llm_changeset(cached, accepted_change_ids, save_as)

    payload: dict[str, Any] = {"changeset_id": changeset_id, "accepted_change_ids": accepted_change_ids}
    if save_as:
        payload["save_as"] = save_as
    return await toolbox_client.call_tool("docmate_apply_changeset", payload)


async def commit_changeset(
    changeset_id: str,
    accepted_change_ids: list[str],
    save_as: str | None = None,
    confirmed_by: str | None = None,
) -> dict[str, Any]:
    """Commit accepted changes and expose the produced file through download."""
    result = await apply_changeset(changeset_id, accepted_change_ids, save_as)
    data = result.get("data") if isinstance(result.get("data"), dict) else {}
    output_path = Path(str(data.get("output_path") or "")) if data.get("output_path") else None
    file_id = _register_download_file(output_path) if output_path and output_path.exists() else ""
    commit_data = {
        **data,
        "status": "committed" if result.get("ok") else "failed",
        "changeset_id": changeset_id,
        "accepted_change_ids": accepted_change_ids,
        "confirmed_by": confirmed_by or "",
    }
    if file_id:
        commit_data["file_id"] = file_id
        commit_data["download_url"] = f"/api/docmate/download/{file_id}"
    return {
        **result,
        "tool": "docmate_commit_changeset",
        "summary": result.get("summary") or ("DocMate 变更已提交。" if result.get("ok") else "DocMate 变更提交失败。"),
        "data": commit_data,
    }


async def retry_changeset(
    changeset_id: str,
    instruction: str | None = None,
    context: Any | None = None,
    feedback: str | None = None,
) -> dict[str, Any]:
    """Regenerate a changeset from the previous doc/instruction with feedback."""
    cached = _CHANGESET_CACHE.get(changeset_id)
    if not cached:
        return {"ok": False, "error": "changeset_not_found", "summary": f"变更集不存在: {changeset_id}"}

    retry_context = {
        "previous_changeset_id": changeset_id,
        "previous_summary": cached.get("summary", {}),
        "feedback": feedback or "",
        "context": context if context is not None else cached.get("context"),
    }
    result = await generate_changeset(
        str(cached.get("doc_id") or ""),
        instruction or str(cached.get("instruction") or ""),
        retry_context,
    )
    if isinstance(result.get("data"), dict):
        result["data"]["retried_from"] = changeset_id
    new_id = _nested(result, "data", "changeset_id")
    if new_id and str(new_id) in _CHANGESET_CACHE:
        _CHANGESET_CACHE[str(new_id)]["retried_from"] = changeset_id
    return result


async def pipeline_edit(file_path: str, instruction: str, save_as: str | None = None, context: Any | None = None) -> dict[str, Any]:
    """Full pipeline: read -> generate -> apply in one call."""
    read_result = await read_docx(file_path)
    if not read_result.get("ok"):
        return {"ok": False, "error": "read_docx failed", "detail": read_result}

    doc_id = str(_nested(read_result, "data", "doc_id") or "")
    if not doc_id:
        return {"ok": False, "error": "read_docx returned no doc_id", "detail": read_result}

    gen_result = await generate_changeset(doc_id, instruction, context)
    if not gen_result.get("ok"):
        return {"ok": False, "error": "generate_changeset failed", "detail": gen_result}

    changeset_id = str(_nested(gen_result, "data", "changeset_id") or "")
    changes = _nested(gen_result, "data", "changes") or []
    accepted = [str(c["change_id"]) for c in changes if isinstance(c, dict) and c.get("change_id")]

    apply_result = await apply_changeset(changeset_id, accepted, save_as)
    if not apply_result.get("ok"):
        return {"ok": False, "error": "apply_changeset failed", "detail": apply_result}

    return {
        "ok": True,
        "steps": {"read": read_result, "generate": gen_result, "apply": apply_result},
        "output_path": _nested(apply_result, "data", "output_path"),
        "backup_path": _nested(apply_result, "data", "backup_path"),
    }


def resolve_download_file(file_id: str) -> Path | None:
    """Resolve a DocMate file id to a local path if it is known and safe."""
    if not re.fullmatch(r"[A-Za-z0-9_-]{8,80}", file_id or ""):
        return None
    cached = _DOWNLOAD_FILES.get(file_id)
    if cached and cached.exists():
        return cached
    for root in (_docmate_upload_dir(), _docmate_output_dir()):
        if not root.exists():
            continue
        for path in root.glob(f"{file_id}_*"):
            if path.is_file():
                _DOWNLOAD_FILES[file_id] = path
                return path
    return None


def _generate_changeset_id() -> str:
    return f"cs_{datetime.now(timezone.utc).strftime('%Y%m%d')}_{uuid.uuid4().hex[:8]}"


async def _generate_changeset_with_llm(doc_id: str, instruction: str, context: Any | None) -> dict[str, Any]:
    if not settings.llm_configured:
        return {"ok": False, "error": "llm_not_configured"}

    try:
        raw = await llm_gateway.complete_document_json(
            _docmate_changeset_system_prompt(),
            _docmate_changeset_user_text(doc_id, instruction, context),
        )
        parsed = _extract_json_object(raw)
        if parsed.get("action") == "clarify":
            questions = parsed.get("clarify_questions") if isinstance(parsed.get("clarify_questions"), list) else []
            return {
                "ok": False,
                "tool": "docmate_generate_changeset",
                "error": "clarification_required",
                "summary": "DocMate 需要补充信息后才能生成修改方案。",
                "data": {"questions": questions},
            }
        if parsed.get("action") != "generate_changes":
            return {"ok": False, "error": f"unsupported_llm_action:{parsed.get('action')}"}
        changes = _normalize_llm_changes(parsed.get("changes"))
        if not changes:
            return {"ok": False, "error": "llm_returned_no_changes"}
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "error": str(exc)}

    changeset_id = _generate_changeset_id()
    preview_cards = [_preview_card_for_change(change, index) for index, change in enumerate(changes, start=1)]
    doc_entry = _DOC_CACHE.get(doc_id, {})
    summary = _changes_summary(changes)
    changeset = {
        "changeset_id": changeset_id,
        "doc_id": doc_id,
        "source_path": doc_entry.get("source_path", ""),
        "instruction": instruction,
        "context": context,
        "changes": changes,
        "preview_cards": preview_cards,
        "summary": summary,
        "generator": "llm",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    _CHANGESET_CACHE[changeset_id] = changeset
    return {
        "ok": True,
        "tool": "docmate_generate_changeset",
        "summary": f"LLM 已生成 {len(preview_cards)} 项修改方案。",
        "data": {
            "changeset_id": changeset_id,
            "doc_id": doc_id,
            "instruction": instruction,
            "changes": changes,
            "preview_cards": preview_cards,
            "total_changes": len(preview_cards),
            "summary": summary,
            "generator": "llm",
        },
    }


def _docmate_changeset_system_prompt() -> str:
    return (
        "You are DocMate, an AI assistant for safe DOCX editing. Return one JSON object only. "
        "Prefer precise, reviewable changes; never claim that changes are committed. "
        "Schema: {\"action\":\"generate_changes\",\"changes\":[{\"change_id\":\"optional\","
        "\"change_type\":\"text_replace|text_insert|text_delete|table_cell_update|table_insert|image_insert|image_replace\","
        "\"target\":{},\"old_content\":\"\",\"new_content\":\"\",\"reason\":\"\",\"risk_level\":\"low|medium|high\","
        "\"confidence\":0.0}]} or {\"action\":\"clarify\",\"clarify_questions\":[\"...\"]}. "
        "Use high risk for compliance or safety meaning changes, medium for substantive wording, low for cosmetic edits."
    )


def _docmate_changeset_user_text(doc_id: str, instruction: str, context: Any | None) -> str:
    payload = {
        "doc_id": doc_id,
        "instruction": instruction,
        "context": context,
        "document_context": _compact_doc_context(_DOC_CACHE.get(doc_id, {})),
    }
    return json.dumps(payload, ensure_ascii=False, default=str)


def _compact_doc_context(doc_entry: dict[str, Any]) -> dict[str, Any]:
    structure = doc_entry.get("structure") if isinstance(doc_entry.get("structure"), dict) else {}
    paragraphs = structure.get("paragraphs") if isinstance(structure.get("paragraphs"), list) else []
    tables = structure.get("tables") if isinstance(structure.get("tables"), list) else []
    return {
        "source_path": doc_entry.get("source_path", ""),
        "paragraphs": paragraphs[:40],
        "tables": tables[:10],
        "image_count": structure.get("image_count", 0),
    }


def _extract_json_object(raw: dict[str, Any]) -> dict[str, Any]:
    if not isinstance(raw, dict):
        raise ValueError("LLM response is not a JSON object")
    if raw.get("available") is False:
        raise ValueError(str(raw.get("reason") or "LLM unavailable"))
    if "action" in raw:
        return raw
    choices = raw.get("choices")
    if isinstance(choices, list) and choices:
        message = choices[0].get("message") if isinstance(choices[0], dict) else {}
        content = message.get("content") if isinstance(message, dict) else ""
        if isinstance(content, str):
            return json.loads(content)
    raise ValueError("model did not return DocMate JSON")


def _normalize_llm_changes(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    changes: list[dict[str, Any]] = []
    for item in value:
        if not isinstance(item, dict):
            continue
        change_type = str(item.get("change_type") or "text_replace")
        risk_level = str(item.get("risk_level") or "medium").lower()
        if change_type not in _VALID_CHANGE_TYPES:
            continue
        if risk_level not in _VALID_RISK_LEVELS:
            risk_level = "medium"
        change = {
            "change_id": str(item.get("change_id") or f"chg_{uuid.uuid4().hex[:8]}"),
            "change_type": change_type,
            "target": item.get("target") if isinstance(item.get("target"), dict) else {},
            "old_content": _stringify_content(item.get("old_content")),
            "new_content": _stringify_content(item.get("new_content")),
            "reason": str(item.get("reason") or ""),
            "risk_level": risk_level,
            "confidence": _bounded_float(item.get("confidence"), default=0.75),
            "status": "pending",
        }
        changes.append(change)
    return changes


def _preview_card_for_change(change: dict[str, Any], index: int) -> dict[str, Any]:
    change_type = str(change.get("change_type") or "text_replace")
    old_content = str(change.get("old_content") or "")
    new_content = str(change.get("new_content") or "")
    return {
        "change_id": change["change_id"],
        "title": f"修改 {index}: {change_type}",
        "old_content": old_content,
        "new_content": new_content,
        "before": old_content,
        "after": new_content,
        "risk_level": change.get("risk_level", "medium"),
        "confidence": change.get("confidence", 0.0),
        "reason": change.get("reason", ""),
        "target": change.get("target", {}),
    }


def _changes_summary(changes: list[dict[str, Any]]) -> dict[str, int]:
    summary = {"total": len(changes), "high": 0, "medium": 0, "low": 0}
    for change in changes:
        risk = str(change.get("risk_level") or "medium")
        if risk in {"high", "medium", "low"}:
            summary[risk] += 1
    return summary


def _apply_llm_changeset(cached: dict[str, Any], accepted_change_ids: list[str], save_as: str | None) -> dict[str, Any]:
    accepted = set(accepted_change_ids)
    changes = [
        change for change in cached.get("changes", [])
        if isinstance(change, dict) and change.get("change_id") in accepted
    ]
    if not changes:
        return {"ok": False, "tool": "docmate_apply_changeset", "error": "no_changes_selected", "summary": "没有选中任何变更，请至少选择一项后再应用。"}

    source_path = Path(str(cached.get("source_path") or ""))
    output_path = Path(save_as) if save_as else _docmate_output_dir() / f"{source_path.stem}_modified_{uuid.uuid4().hex[:6]}.docx"
    backup_path = output_path.parent / f"{output_path.stem}_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        from docx import Document

        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_path, backup_path)
        doc = Document(str(source_path))
        applied = 0
        errors: list[str] = []
        for change in changes:
            ok = _apply_change_to_document(doc, change)
            if ok:
                applied += 1
            else:
                errors.append(f"未能定位变更 {change.get('change_id')} 的原文。")
        doc.save(str(output_path))
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "tool": "docmate_apply_changeset", "error": str(exc), "summary": f"保存文档失败: {exc}"}

    return {
        "ok": True,
        "tool": "docmate_apply_changeset",
        "summary": f"已应用 {applied} 项修改。",
        "data": {
            "applied": applied,
            "rejected": max(0, len(cached.get("changes", [])) - len(changes)),
            "errors": errors,
            "output_path": str(output_path),
            "backup_path": str(backup_path),
            "source_path": str(source_path),
        },
    }


def _apply_change_to_document(doc: Any, change: dict[str, Any]) -> bool:
    change_type = str(change.get("change_type") or "text_replace")
    old_content = str(change.get("old_content") or "")
    new_content = str(change.get("new_content") or "")
    if change_type == "text_insert":
        doc.add_paragraph(new_content)
        return True
    if change_type == "text_delete":
        new_content = ""
    if change_type not in {"text_replace", "text_delete", "table_cell_update"}:
        return False
    if not old_content:
        return False
    return _replace_in_paragraphs(doc.paragraphs, old_content, new_content) or _replace_in_tables(doc.tables, old_content, new_content)


def _replace_in_tables(tables: Any, old_content: str, new_content: str) -> bool:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                if _replace_in_paragraphs(cell.paragraphs, old_content, new_content):
                    return True
    return False


def _replace_in_paragraphs(paragraphs: Any, old_content: str, new_content: str) -> bool:
    for paragraph in paragraphs:
        if old_content not in paragraph.text:
            continue
        for run in paragraph.runs:
            if old_content in run.text:
                run.text = run.text.replace(old_content, new_content, 1)
                return True
        paragraph.text = paragraph.text.replace(old_content, new_content, 1)
        return True
    return False


def _remember_doc(file_path: str, result: dict[str, Any]) -> None:
    if not result.get("ok"):
        return
    data = result.get("data") if isinstance(result.get("data"), dict) else {}
    doc_id = str(data.get("doc_id") or result.get("doc_id") or "")
    if not doc_id:
        return
    structure = data.get("structure") if isinstance(data.get("structure"), dict) else data
    source_path = str(data.get("source_path") or result.get("source_path") or file_path)
    _DOC_CACHE[doc_id] = {"doc_id": doc_id, "source_path": source_path, "structure": structure, "raw": result}


def _remember_toolbox_changeset(result: dict[str, Any], doc_id: str, instruction: str, context: Any | None) -> None:
    if not result.get("ok"):
        return
    data = result.get("data") if isinstance(result.get("data"), dict) else {}
    changeset_id = str(data.get("changeset_id") or "")
    if not changeset_id:
        return
    _CHANGESET_CACHE[changeset_id] = {
        "changeset_id": changeset_id,
        "doc_id": doc_id,
        "source_path": _DOC_CACHE.get(doc_id, {}).get("source_path", ""),
        "instruction": instruction,
        "context": context,
        "changes": data.get("changes") if isinstance(data.get("changes"), list) else [],
        "preview_cards": data.get("preview_cards") if isinstance(data.get("preview_cards"), list) else [],
        "summary": data.get("summary") if isinstance(data.get("summary"), dict) else {},
        "generator": "toolbox",
    }


def _register_download_file(path: Path | None, *, file_id: str | None = None) -> str:
    if not path:
        return ""
    resolved = Path(path)
    key = file_id or uuid.uuid4().hex
    _DOWNLOAD_FILES[key] = resolved
    return key


def _docmate_upload_dir() -> Path:
    return settings.chitung_data_dir / "docmate_uploads"


def _docmate_output_dir() -> Path:
    return settings.chitung_data_dir / "docmate_outputs"


def _safe_filename(filename: str) -> str:
    name = Path(filename).name.strip() or "document.docx"
    return re.sub(r"[^A-Za-z0-9._ -]+", "_", name)


def _stringify_content(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    return json.dumps(value, ensure_ascii=False, default=str)


def _bounded_float(value: Any, *, default: float) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError):
        number = default
    return max(0.0, min(number, 1.0))


def _nested(value: dict[str, Any], *keys: str) -> Any:
    current: Any = value
    for key in keys:
        if not isinstance(current, dict):
            return None
        current = current.get(key)
    return current
