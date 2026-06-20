"""
docx_executor.py — 执行 ChangeSet 中的变更到 .docx 文件

按《DocMate EHS docx升级与Skill落地实施方案 v1》第 9.3 节伪代码实现。
支持：text_replace, text_insert, text_delete, table_cell_update,
      table_insert, image_insert, image_replace
"""

from __future__ import annotations

import copy
import os
import re
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .docx_model import (
    ChangeType,
    ChangeStatus,
    RiskLevel,
    DocumentChange,
    ChangeSet,
    DocumentStructure,
    make_hash_id,
)


# ─── 1. 核心执行函数 ────────────────────────────────────────────

def apply_changeset(
    changeset: ChangeSet,
    output_path: str,
    structure: Optional[DocumentStructure] = None,
) -> dict:
    """
    主入口：将已接受的变更应用到 .docx 文件。
    对应方案第 9.3 节。

    Args:
        changeset: 包含 accepted 变更的 ChangeSet
        output_path: 输出路径
        structure: 可选，文档结构（用于锚点定位）

    Returns:
        {"ok": True/False, "data": {"output_path": ..., "applied": N, "errors": [...]}}
    """
    source_path = changeset.source_path
    if not source_path or not os.path.exists(source_path):
        return {
            "ok": False,
            "error": f"Source file not found: {source_path}",
            "error_code": "DOCX_001_PARSE_FAILED",
        }

    try:
        doc = DocxDocument(source_path)
    except Exception as e:
        return {
            "ok": False,
            "error": f"Failed to open document: {e}",
            "error_code": "DOCX_001_PARSE_FAILED",
        }

    accepted = changeset.get_accepted_changes()
    applied = 0
    errors = []

    for chg in accepted:
        try:
            _apply_single_change(doc, chg)
            chg.status = ChangeStatus.ACCEPTED
            applied += 1
        except Exception as e:
            errors.append({
                "change_id": chg.change_id,
                "error": str(e),
                "error_code": "DOCX_003_APPLY_FAILED",
            })

    try:
        output_path = str(Path(output_path).resolve())
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
    except Exception as e:
        return {
            "ok": False,
            "error": f"Failed to save document: {e}",
            "error_code": "DOCX_003_APPLY_FAILED",
        }

    return {
        "ok": True,
        "data": {
            "output_path": output_path,
            "applied": applied,
            "rejected": len(accepted) - applied,
            "errors": errors,
        },
    }


def _apply_single_change(doc: DocxDocument, chg: DocumentChange):
    """路由单条变更到对应处理器。"""
    ct = chg.change_type

    if ct == ChangeType.TEXT_REPLACE:
        _text_replace(doc, chg)
    elif ct == ChangeType.TEXT_INSERT:
        _text_insert(doc, chg)
    elif ct == ChangeType.TEXT_DELETE:
        _text_delete(doc, chg)
    elif ct == ChangeType.TABLE_CELL_UPDATE:
        _table_cell_update(doc, chg)
    elif ct == ChangeType.TABLE_INSERT:
        _table_insert(doc, chg)
    elif ct == ChangeType.IMAGE_INSERT:
        _image_insert(doc, chg)
    elif ct == ChangeType.IMAGE_REPLACE:
        _image_replace(doc, chg)
    else:
        raise ValueError(f"Unknown change_type: {ct}")


# ─── 2. 文本修改 ────────────────────────────────────────────────

def _find_paragraph_by_id(doc: DocxDocument, paragraph_id: str) -> Optional[tuple]:
    """根据 paragraph_id (如 P12#a13f) 找到段落及其索引。"""
    # 解析 prefix + index
    match = re.match(r"([A-Z]+)(\d+)#", paragraph_id)
    if not match:
        return None
    prefix = match.group(1)
    target_index = int(match.group(2)) - 1  # 转为 0-based

    if prefix == "P":
        if 0 <= target_index < len(doc.paragraphs):
            return (doc.paragraphs[target_index], target_index)
    return None


def _text_replace(doc: DocxDocument, chg: DocumentChange):
    """替换段落中的文字。"""
    pid = chg.target.paragraph_id
    if not pid:
        raise ValueError("text_replace requires paragraph_id in target")

    result = _find_paragraph_by_id(doc, pid)
    if not result:
        raise ValueError(f"Paragraph not found: {pid} (DOCX_002_ANCHOR_NOT_FOUND)")

    para, idx = result
    old_text = chg.old_content
    new_text = chg.new_content

    # 策略：在段落 runs 中查找 old_text 并替换
    full_text = para.text
    if old_text and old_text in full_text:
        # 精确替换：inline 替换 run 文本
        _replace_in_runs(para, old_text, new_text)
    else:
        # 全局替换整段
        _replace_paragraph_text(para, new_text)


def _replace_in_runs(para, old_text: str, new_text: str):
    """在段落 runs 中 inline 替换文字，保持格式。"""
    remaining_old = old_text
    remaining_new = new_text

    for run in para.runs:
        if not remaining_old:
            break
        if remaining_old in run.text:
            run.text = run.text.replace(remaining_old, remaining_new, 1)
            break
        elif run.text in remaining_old:
            # run 是 old_text 的一部分
            run.text = remaining_new[:len(run.text)] if len(remaining_new) >= len(run.text) else ""
            remaining_old = remaining_old[len(run.text):]
            remaining_new = remaining_new[len(run.text):]
        else:
            # 部分匹配
            common_prefix = _common_prefix(run.text, remaining_old)
            if common_prefix:
                run.text = remaining_new[:len(common_prefix)]
                remaining_old = remaining_old[len(common_prefix):]
                remaining_new = remaining_new[len(common_prefix):]


def _common_prefix(a: str, b: str) -> str:
    """计算两个字符串的公共前缀。"""
    i = 0
    for i, (ca, cb) in enumerate(zip(a, b)):
        if ca != cb:
            break
    else:
        i = min(len(a), len(b))
    return a[:i]


def _replace_paragraph_text(para, new_text: str):
    """替换段落全部文字，保留第一个 run 的格式。"""
    if para.runs:
        # 保留第一个 run 的格式
        first_run = para.runs[0]
        # 清除后续 runs
        for run in para.runs[1:]:
            run.text = ""
        first_run.text = new_text
    else:
        para.add_run(new_text)


def _text_insert(doc: DocxDocument, chg: DocumentChange):
    """在指定段落之后插入新段落。"""
    pid = chg.target.paragraph_id
    if not pid:
        raise ValueError("text_insert requires paragraph_id in target")

    result = _find_paragraph_by_id(doc, pid)
    if not result:
        raise ValueError(f"Paragraph not found: {pid}")

    para, idx = result
    # 在 para 的 XML 元素后插入新段落
    new_para = doc.add_paragraph(chg.new_content)

    # 继承原段落格式
    if para.style:
        new_para.style = para.style

    # 移动新段落到原段落之后
    para._element.addnext(new_para._element)


def _text_delete(doc: DocxDocument, chg: DocumentChange):
    """删除指定段落。"""
    pid = chg.target.paragraph_id
    if not pid:
        raise ValueError("text_delete requires paragraph_id in target")

    result = _find_paragraph_by_id(doc, pid)
    if not result:
        raise ValueError(f"Paragraph not found: {pid}")

    para, idx = result
    # 移除段落的 XML 元素
    para._element.getparent().remove(para._element)


# ─── 3. 表格操作 ────────────────────────────────────────────────

def _find_table_by_id(doc: DocxDocument, table_id: str) -> Optional[tuple]:
    """根据 table_id (如 T2#...) 找到表格。"""
    match = re.match(r"T(\d+)#", table_id)
    if not match:
        return None
    target_index = int(match.group(1)) - 1
    if 0 <= target_index < len(doc.tables):
        return (doc.tables[target_index], target_index)
    return None


def _table_cell_update(doc: DocxDocument, chg: DocumentChange):
    """更新表格单元格内容。"""
    tid = chg.target.table_id
    if not tid:
        raise ValueError("table_cell_update requires table_id in target")

    result = _find_table_by_id(doc, tid)
    if not result:
        raise ValueError(f"Table not found: {tid} (DOCX_002_ANCHOR_NOT_FOUND)")

    table, t_idx = result
    cell_ref = chg.target.cell
    if not cell_ref:
        raise ValueError("table_cell_update requires cell in target")

    row_idx = cell_ref.get("row", 1) - 1
    col_idx = cell_ref.get("col", 1) - 1

    if row_idx >= len(table.rows):
        raise ValueError(f"Row {row_idx + 1} out of range (table has {len(table.rows)} rows)")
    if col_idx >= len(table.columns):
        raise ValueError(f"Column {col_idx + 1} out of range (table has {len(table.columns)} cols)")

    cell = table.cell(row_idx, col_idx)
    # 清除并替换单元格文字
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].add_run(chg.new_content)
    else:
        cell.add_paragraph().add_run(chg.new_content)


def _table_insert(doc: DocxDocument, chg: DocumentChange):
    """在指定段落之后插入表格。"""
    pid = chg.target.paragraph_id
    if not pid:
        raise ValueError("table_insert requires paragraph_id in target (as anchor)")

    result = _find_paragraph_by_id(doc, pid)
    if not result:
        raise ValueError(f"Anchor paragraph not found: {pid}")

    para, idx = result
    rows = chg.table_rows or 3
    cols = chg.table_cols or 3
    data = chg.table_data or []

    table = doc.add_table(rows=rows, cols=cols, style="Table Grid")

    # 填充数据
    for r in range(min(rows, len(data))):
        for c in range(min(cols, len(data[r]) if r < len(data) else 0)):
            if r < len(data) and c < len(data[r]):
                table.cell(r, c).text = data[r][c]

    # 移动表格到锚点段落之后
    para._element.addnext(table._element)


# ─── 4. 图片操作 ────────────────────────────────────────────────

def _image_insert(doc: DocxDocument, chg: DocumentChange):
    """在指定段落之后插入图片（从本地文件夹）。"""
    pid = chg.target.anchor or chg.target.paragraph_id
    if not pid:
        raise ValueError("image_insert requires anchor or paragraph_id in target")

    result = _find_paragraph_by_id(doc, pid)
    if not result:
        raise ValueError(f"Anchor paragraph not found: {pid}")

    para, idx = result
    image_path = chg.image_path

    if not image_path or not os.path.exists(image_path):
        raise ValueError(f"Image file not found: {image_path}")

    # 插入图片
    width_inches = chg.image_width_px / 96.0 if chg.image_width_px else 5.5
    new_para = doc.add_paragraph()
    run = new_para.add_run()
    run.add_picture(image_path, width=Inches(min(width_inches, 6.0)))

    # 移动图片段落到锚点段落之后
    para._element.addnext(new_para._element)


def _image_replace(doc: DocxDocument, chg: DocumentChange):
    """替换文档中已有的图片（通过删除原图片段落 + 插入新图片实现）。"""
    pid = chg.target.anchor or chg.target.paragraph_id
    if not pid:
        raise ValueError("image_replace requires anchor in target")

    # 1. 删除原图片段落
    result = _find_paragraph_by_id(doc, pid)
    if result:
        old_para, idx = result
        old_para._element.getparent().remove(old_para._element)

    # 2. 在当前锚点之前找到合适位置插入
    # 复用 _image_insert 的逻辑
    chg.change_type = ChangeType.IMAGE_INSERT  # 临时切换
    _image_insert(doc, chg)


# ─── 5. 变更回滚 ────────────────────────────────────────────────

def revert_changeset(
    changeset: ChangeSet,
    backup_path: str,
    output_path: str,
) -> dict:
    """
    回滚变更：直接将备份文件复制回目标路径。
    注意：更精细的回滚需要保存原始 snapshot，这里用文件级回滚。
    """
    if not os.path.exists(backup_path):
        return {"ok": False, "error": f"Backup not found: {backup_path}"}

    try:
        import shutil
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        shutil.copy2(backup_path, output_path)
        for chg in changeset.changes:
            if chg.status == ChangeStatus.ACCEPTED:
                chg.status = ChangeStatus.REVERTED
        return {"ok": True, "data": {"output_path": output_path}}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── 6. 备份 ────────────────────────────────────────────────────

def create_backup(source_path: str) -> str:
    """创建 .docx 文件的备份副本。"""
    backup_path = source_path + ".backup"
    import shutil
    shutil.copy2(source_path, backup_path)
    return backup_path
