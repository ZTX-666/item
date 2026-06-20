import os
import shutil
import unicodedata
import json
from datetime import datetime
from pathlib import Path
from typing import List

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image

from models import (
    Photo,
    PhotoClassification,
    Audio,
    AudioTranscription,
    GenerateDiarioRequest,
    PhotoOrderRequest,
    UploadResponse,
    DiarioResponse,
)


# Funções auxiliares
def sanitize_filename(text: str) -> str:
    """
    Sanitiza texto para uso seguro em nomes de arquivos.
    Remove acentos, caracteres especiais e espaços.
    """
    # Normaliza para NFD (decompõe caracteres acentuados)
    nfkd = unicodedata.normalize('NFD', text)
    # Remove acentos (mantém apenas ASCII)
    ascii_text = nfkd.encode('ascii', 'ignore').decode('ascii')
    # Remove caracteres especiais (mantém alfanuméricos, espaços e underscores)
    safe_text = ''.join(c if c.isalnum() or c in [' ', '_', '-'] else '_' for c in ascii_text)
    # Substitui espaços por underscores
    return safe_text.replace(' ', '_')


# Configurações
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
AUDIO_DIR = Path("uploads/audio")
AUDIO_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

# Carregar variáveis de ambiente
load_dotenv()

# Configurar Gemini AI
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    # Usar Gemini 2.0 Flash (FREE tier - 1500 req/dia)
    gemini_vision_model = genai.GenerativeModel('gemini-2.0-flash-exp')
    gemini_text_model = genai.GenerativeModel('gemini-2.0-flash-exp')
else:
    gemini_vision_model = None
    gemini_text_model = None
    print("⚠️  AVISO: GEMINI_API_KEY não configurada. Usando modo MOCK.")

# Memória temporária das fotos e áudios (em produção, usar PostgreSQL)
photos_storage: dict = {}
audio_storage: dict = {}

app = FastAPI(title="Diário de Obras API", version="1.0.0")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
async def root():
    return {"message": "Diário de Obras API v1.0.0"}


@app.post("/api/fotos/upload", response_model=UploadResponse)
async def upload_foto(file: UploadFile = File(...)):
    """Upload de uma foto e retorna URL"""
    try:
        # Validar tipo de arquivo
        if not file.content_type or not file.content_type.startswith("image/"):
            raise HTTPException(status_code=400, detail="Apenas imagens são permitidas")

        # Gerar nome único
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = file.filename if file.filename else "image.jpg"
        filename_safe = filename.replace(".", "_")
        photo_id = f"photo_{timestamp}_{filename_safe}"
        file_path = UPLOAD_DIR / f"{timestamp}_{filename}"

        # Salvar arquivo
        with file_path.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Armazenar metadados
        photos_storage[photo_id] = {
            "id": photo_id,
            "name": file.filename,
            "url": f"http://localhost:8000/api/fotos/{timestamp}_{filename}",
            "path": str(file_path),
            "size": file.size if file.size else 0,
            "type": file.content_type or "image/jpeg",
        }

        return UploadResponse(
            success=True, message="Foto carregada com sucesso", photo_id=photo_id
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao carregar foto: {str(e)}")


@app.get("/api/fotos/{filename}")
async def get_foto(filename: str):
    """Retorna a foto (usada pelo frontend para preview)"""
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Foto não encontrada")
    return FileResponse(file_path)


@app.post("/api/fotos/classificar")
async def classificar_fotos(photo_ids: List[str]):
    """Classificar fotos usando Gemini Vision"""
    classifications = []

    for photo_id in photo_ids:
        if photo_id not in photos_storage:
            continue

        try:
            # Carregar imagem
            photo_path = Path(photos_storage[photo_id]["path"])
            if not photo_path.exists():
                continue

            # Usar Gemini Vision se disponível
            if gemini_vision_model:
                # Abrir imagem
                img = Image.open(photo_path)

                # Prompt para classificação de fotos de obra
                prompt = """Analise esta foto de construção civil e forneça:

1. **Descrição**: Descreva o que vê na foto (máx. 2 frases)
2. **Categoria**: Classifique em UMA das categorias:
   - estrutura (concreto, vigas, pilares, lajes)
   - alvenaria (paredes, blocos, tijolos)
   - revestimento (reboco, pintura, acabamentos)
   - instalacoes (elétrica, hidráulica, AVAC)
   - acabamento (pisos, azulejos, portas, janelas)
   - seguranca (EPIs, sinalização, proteções)
   - equipamentos (máquinas, ferramentas)
   - outros (geral da obra)

3. **Tags**: Liste 3-5 palavras-chave relevantes

Retorne APENAS no formato JSON:
{"description": "...", "category": "...", "tags": ["...", "..."]}"""

                # Gerar classificação
                response = gemini_vision_model.generate_content([prompt, img])
                result_text = response.text.strip()

                # Parsear JSON da resposta
                # Remover markdown se presente
                if result_text.startswith("```json"):
                    result_text = result_text.split("```json")[1].split("```")[0].strip()
                elif result_text.startswith("```"):
                    result_text = result_text.split("```")[1].split("```")[0].strip()

                result = json.loads(result_text)

                classification = PhotoClassification(
                    photo_id=photo_id,
                    description=result.get("description", "Foto de obra"),
                    category=result.get("category", "outros"),
                    tags=result.get("tags", ["obra"]),
                    confidence=0.9,  # Gemini Vision tem alta confiança
                )
            else:
                # Fallback para mock se Gemini não disponível
                classification = PhotoClassification(
                    photo_id=photo_id,
                    description="Foto de obra em andamento (classificação mock)",
                    category="estrutura",
                    tags=["obra", "construção"],
                    confidence=0.8,
                )

            classifications.append(classification)

        except Exception as e:
            # Em caso de erro, retorna classificação genérica
            print(f"Erro ao classificar {photo_id}: {str(e)}")
            classification = PhotoClassification(
                photo_id=photo_id,
                description="Foto de obra (erro na classificação)",
                category="outros",
                tags=["obra"],
                confidence=0.5,
            )
            classifications.append(classification)

    return classifications


@app.post("/api/audio/upload", response_model=UploadResponse)
async def upload_audio(file: UploadFile = File(...)):
    """Upload de áudio e retorna URL"""
    try:
        # Validar tipo de arquivo
        if not file.content_type or not file.content_type.startswith("audio/"):
            raise HTTPException(status_code=400, detail="Apenas áudios são permitidos")

        # Gerar nome único
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = file.filename if file.filename else "audio.mp3"
        filename_safe = filename.replace(".", "_")
        audio_id = f"audio_{timestamp}_{filename_safe}"
        file_path = AUDIO_DIR / f"{timestamp}_{filename}"

        # Salvar arquivo
        with file_path.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Armazenar metadados
        audio_storage[audio_id] = {
            "id": audio_id,
            "name": file.filename,
            "url": f"http://localhost:8000/api/audio/{timestamp}_{filename}",
            "path": str(file_path),
            "size": file.size if file.size else 0,
            "type": file.content_type or "audio/mpeg",
        }

        return UploadResponse(
            success=True, message="Áudio carregado com sucesso", audio_id=audio_id
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao carregar áudio: {str(e)}")


@app.get("/api/audio/{filename}")
async def get_audio(filename: str):
    """Retorna o áudio"""
    file_path = AUDIO_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Áudio não encontrado")
    return FileResponse(file_path)


@app.post("/api/audio/transcrever")
async def transcrever_audio(audio_id: str):
    """Transcrever áudio usando Gemini STT"""
    if audio_id not in audio_storage:
        raise HTTPException(status_code=404, detail="Áudio não encontrado")

    try:
        audio_path = Path(audio_storage[audio_id]["path"])
        if not audio_path.exists():
            raise HTTPException(status_code=404, detail="Arquivo de áudio não encontrado")

        # Usar Gemini STT se disponível
        if gemini_text_model:
            # Upload do arquivo de áudio para Gemini
            audio_file = genai.upload_file(path=str(audio_path))

            # Prompt para transcrição contextualizada
            prompt = """Transcreva este áudio de um engenheiro relatando atividades de obra de construção civil.

Forneça uma transcrição limpa e organizada do que foi falado, mantendo:
- Fidelidade ao conteúdo (não invente informações)
- Pontuação adequada
- Paragrafação para facilitar leitura

Retorne APENAS a transcrição, sem introduções ou comentários."""

            # Gerar transcrição
            response = gemini_text_model.generate_content([prompt, audio_file])
            transcription = response.text.strip()

            # Limpar o arquivo temporário do Gemini
            genai.delete_file(audio_file.name)

        else:
            # Fallback para mock se Gemini não disponível
            transcription = f"""Transcrição do áudio: {audio_storage[audio_id]["name"]}

NOTA: Gemini API não configurada. Esta é uma transcrição de exemplo.

Atividades realizadas hoje:
- Instalação de vergalhões no segundo pavimento
- Pintura de paredes externas (50% concluído)
- Revisão de instalações elétricas
- Reunião com equipe de segurança"""

        return {
            "audio_id": audio_id,
            "transcription": transcription,
            "language": "pt",
            "duration": 0,  # Gemini não retorna duração
        }

    except Exception as e:
        # Em caso de erro, retorna mensagem de erro
        print(f"Erro ao transcrever áudio {audio_id}: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao transcrever áudio: {str(e)}"
        )


@app.post("/api/diario/gerar", response_model=DiarioResponse)
async def gerar_diario(request: GenerateDiarioRequest):
    """Gerar Diário de Obra em DOCX"""
    try:
        # Importar python-docx
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Criar documento
        doc = Document()

        # Título
        title = doc.add_heading("DIÁRIO DE OBRA", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informações do projeto
        doc.add_heading("Informações do Projeto", 1)
        info = doc.add_paragraph()
        info.add_run(f"Nome do Projeto: {request.project_name}\n")
        info.add_run(f"Local: {request.project_location}\n")
        if request.contractor:
            info.add_run(f"Contratada: {request.contractor}\n")
        if request.supervisor:
            info.add_run(f"Responsável: {request.supervisor}\n")

        # Data
        doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")

        # Fotos classificadas
        doc.add_heading("Fotos e Descrições", 1)

        for photo in request.photos:
            doc.add_heading(f"{photo.category.upper()}", 2)

            # Descrição
            p = doc.add_paragraph()
            p.add_run(f"Descrição: {photo.description}\n")
            p.add_run(f"Tags: {', '.join(photo.tags)}\n")

            # Foto do storage
            if photo.photo_id in photos_storage:
                photo_path = Path(photos_storage[photo.photo_id]["path"])
                if photo_path.exists():
                    doc.add_picture(str(photo_path), width=Inches(6))

            doc.add_paragraph()  # Espaçamento

        # Transcrição de áudio (se houver)
        if request.audio_transcription:
            doc.add_heading("Transcrição de Áudio", 1)
            audio_para = doc.add_paragraph()
            audio_para.add_run(request.audio_transcription)

        # Salvar documento
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = sanitize_filename(request.project_name)
        filename = f"diario_obra_{safe_name}_{timestamp}.docx"
        file_path = OUTPUT_DIR / filename
        doc.save(str(file_path))

        return DiarioResponse(
            success=True,
            message="Diário de Obra gerado com sucesso",
            download_url=f"http://localhost:8000/api/diario/download/{filename}",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao gerar diário: {str(e)}")


@app.get("/api/diario/download/{filename}")
async def download_diario(filename: str):
    """Download do diário gerado"""
    file_path = OUTPUT_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Diário não encontrado")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
