# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

desktop = Path(r"C:\Users\LENOVO\Desktop")
out_path = desktop / "赤瞳安全智能平台-大赛评审问答手册.docx"

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=20, bold=True, color=(23, 58, 130))


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=11, color=(89, 105, 129))


def add_h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=14, bold=True, color=(23, 58, 130))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True, color=(35, 100, 196))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_q(text):
    p = doc.add_paragraph()
    r = p.add_run("问：" + text)
    set_run_font(r, size=11, bold=True)
    p.paragraph_format.space_before = Pt(8)


def add_a(text):
    p = doc.add_paragraph()
    r = p.add_run("答：" + text)
    set_run_font(r, size=11)
    p.paragraph_format.space_after = Pt(4)


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, size=10.5)


add_title("赤瞳安全智能平台")
add_subtitle("中海集团 AI 大赛暨人文中海大赛 · 评审问答手册（带参考答案）")
add_subtitle("产品定位：平台赋能、人机协同、现场共创的安全管理 AI 原生平台")
doc.add_paragraph()

add_h1("一、汇报四项核心内容（评委打分主线）")

add_h2("① 真实业务痛点与应用场景")
add_bullet(
    "痛点：香港工地安全管理中，摄像头、WhatsApp 群、Excel 台账、公文表格、飞书通知长期分散运行；"
    "隐患发现、制度查询、整改通知、外部风险简报大量依赖人工看屏、翻群、填表、催办。"
)
add_bullet(
    "场景：视觉巡检发现临边/未戴安全帽 → 制度依据检索 → 整改通知生成 → 飞书/WhatsApp 确认发送 → 复查闭环；"
    "群消息 @赤瞳 问答与隐患 intake；每日外部风险简报自动生成。"
)
add_bullet("适用对象：地盘安全主任、安健环人员、内派同事及工程管理团队。")

add_h2("② AI 落地实现的技术实施过程")
add_bullet(
    "采用“三层架构”：Electron + Vue 3 桌面工作台 → FastAPI 赤瞳中台（Intent Router、Workflow Engine、"
    "LLM Gateway、RAG、人审确认）→ AgentToolbox 工具层（VLM/OCR/DOCX/飞书/WhatsApp/SQLite）。"
)
add_bullet(
    "落地路径：先打通 1 条隐患闭环 + 3 个演示亮点（视觉巡检、智能填表、外部风险简报），"
    "再按 Skill / Workflow 模块化复制到其他场景。"
)
add_bullet(
    "关键原则：AI 只生成草稿、卡片和候选结论；外发、写文件、关闭隐患等动作必须人工确认，全链路审计留痕。"
)

add_h2("③ 经济价值与管理价值")
add_bullet(
    "管理价值：把“最懂现场的人定义问题、最懂业务的人驱动优化”落到系统里，"
    "推动数字化从“IT 集中开发”走向“平台赋能、现场共创”。"
)
add_bullet(
    "经济价值（预估，试点验证中）：减少安全员重复填表、翻群归档、手工写通知时间；"
    "提升隐患响应与简报时效；降低信息在多层传递中的失真。"
)
add_bullet("组织价值：内派同事可在中台上带领地盘同事自行组装 Skill 和工作流，缩短“需求—上线”周期。")

add_h2("④ 可复制、规模化推广前景")
add_bullet("可复制单元不是整包系统，而是“中台底座 + Skill 包 + Workflow 模板 + 确认策略”。")
add_bullet(
    "同类场景：不同地盘可复用 hazard-intake、visual-patrol、daily-risk-briefing、knowledge-query、"
    "docmate-edit 等能力，仅调整制度库、摄像头源和通知渠道。"
)
add_bullet("推广节奏：单地盘试点 → 同类型项目复制 → 集团内管理增效赛道横向推广；与中海通机器人、飞书生态协同。")

add_h1("二、痛点与场景类问答")

add_q("这个业务痛点目前是怎么解决的？如果不做 AI，传统方案的成本和效率如何？")
add_a(
    "目前主要靠“人 + 群 + 表 + 摄像头”组合：安全员看 CCTV 或巡场拍照，在 WhatsApp/飞书群里通报，"
    "再手工填写 Excel 台账、Word 整改通知和各类安全表格，制度依据靠经验或人工翻文档。"
    "传统方案并非无效，但存在三个结构性问题：一是信息分散在多个工具里，汇总耗时；"
    "二是需求与文档从地盘传到 IT/总部再开发系统，周期长、场景变化快；"
    "三是优秀做法难以沉淀为可复用流程。"
    "粗略估计，安全员每周有大量时间耗在重复归档、填表和通知起草上，且夜间/周末外部风险信息容易滞后。"
)

add_q("这个痛点是普遍存在的共性问题，还是特定场景下的偶发问题？")
add_a(
    "这是工程安全管理中的共性问题，在香港高密度、多分包、强监管工地尤为突出。"
    "群消息多、表格多、摄像头多、制度文档多，是几乎所有大型项目的日常状态。"
    "偶发的是具体隐患类型，共性的是“发现—依据—通知—跟踪—复查”链路上的信息分散与人工串联。"
    "赤瞳针对的是这类共性底座问题，而非单点偶发事件。"
)

add_q("为什么选择用 AI 方案来解决这个问题，AI 的不可替代性体现在哪里？")
add_a(
    "三类环节 AI 具有明显增益，而传统规则系统难以独立完成："
    "① 非结构化理解——从群聊、图片、公文、制度 PDF 中抽取语义；"
    "② 多源融合——把视觉检测结果、RAG 制度条款、模板字段预填、外部新闻信号组合成可执行草稿；"
    "③ 自然语言交互——现场人员用口语描述场景即可触发工作流，无需等待 IT 开发新菜单。"
    "但 AI 不可替代人工最终判断：我们坚持 Human-in-the-loop，所有高风险动作进入待确认队列。"
    "AI 的价值是“提效与辅助决策”，不是“替代管理责任”。"
)

add_h1("三、技术与实施类问答")

add_q("项目落地过程中，遇到的最大技术难点或业务适配难点是什么？最终是如何解决的？")
add_a(
    "主要难点有四类："
    "① 现场工具割裂——已通过 AgentToolbox 统一封装 VLM、OCR、DOCX、飞书、WhatsApp、SQLite 等工具，中台只编排、不重复造轮子；"
    "② 香港场景适配——制度库、繁体/英文混排文档、外部天气与监管信息源单独配置；"
    "③ AI 不可控——采用规则优先 + LLM fallback 的混合路由，且外发/写文件/关案必须人工确认；"
    "④ 端到端闭环尚未完全产品化——当前策略是先做可演示 MVP，优先跑通“隐患 intake、视觉候选、简报草稿、文档 Diff 预览”四段，"
    "再逐步串成完整状态机。"
)

add_q("方案用到的核心数据来源是什么？数据是否经过脱敏、合规校验？是否存在数据安全或知识产权相关风险？")
add_a(
    "数据来源包括：本地 CCTV/RTMP 截图、项目制度与表格模板、本地 WhatsApp wacli 归档库、"
    "公开或授权的外部天气/安全资讯、用户上传 DOCX/PDF 及现场图片。"
    "平台默认本地优先：敏感台账、聊天归档、案件记录优先存本地 SQLite；"
    "调用云端 LLM 前经 LLM Gateway 做上下文脱敏与压缩，仅传必要片段。"
    "知识产权方面，动态扩展限定为 SKILL.md 指令包，不允许 Skill 直接执行脚本或越权访问；"
    "制度文档和模板按项目授权使用。"
    "风险点已正视：云 LLM 调用、群消息内容和影像数据需按集团信息安全制度分级管理，"
    "生产部署应配置专用模型通道、权限审计和数据出境评估。"
)

add_h1("四、价值与成效类问答")

add_q('这个价值是“预估”的还是“已经验证”的？验证数据来源是什么？')
add_a(
    "当前属于“试点验证 + 部分场景已演示”阶段。"
    "已验证能力包括：视觉检测生成候选隐患、RAG 检索制度条款、DocMate 文档 Diff 预览、"
    "外部风险抓取与简报草稿、飞书/WhatsApp 事件接入与中台路由、待确认事项与审计日志。"
    "预估价值主要来自安全员访谈与时间分解：重复填表、通知起草、群消息归档、制度查找等环节可显著压缩。"
    "下一步会在单地盘试点记录 before/after 工时、隐患响应时长和简报发布时效，形成量化对比。"
    "汇报建议诚实表述为“能力已验证，效益待试点量化”。"
)

add_q("为了实现当前效果，项目投入的人力、算力、资金成本大概是多少？投入产出比如何？")
add_a(
    "当前为研发型 MVP：核心团队以小规模敏捷组为主（产品/内派业务专家 + 后端 + 前端 + 工具集成），"
    "复用开源与现有中海通/飞书/Wacli 能力，未大规模采购专有硬件。"
    "算力方面，视觉与 OCR 可本地运行，LLM 按 token 调用，成本可控。"
    "相比传统“每个场景单独立项开发系统”，平台化路线的边际成本更低——"
    "新增一个 Skill 或 Workflow 主要是配置与验证，而非从零开发。"
    "投入产出比的关键在于：底座建设一次性投入，多地盘复用后单次场景成本快速下降。"
)

add_q("有没有对标过行业同类项目的价值产出？你们的优势或差距在哪里？")
add_a(
    "行业常见方案分三类：纯 CCTV 告警平台、纯文档/OCR 工具、纯聊天机器人。"
    "赤瞳差异在于“安全场景一体化编排”——视觉、通讯、制度、文档、确认审批在同一中台闭环，且面向香港工地本地化。"
    "优势：人机协同治理清晰、本地优先、与 WhatsApp/飞书/中海通衔接、支持现场共创 Skill。"
    "差距：Workflow 完整度、飞书端到端确认闭环、视觉策略自适应和 32 Skill 规划仍处 MVP 阶段，"
    "尚未达到全集团规模化生产级别。"
    "我们优势不在“单点算法最强”，而在“更贴近现场工作流的整体落地性”。"
)

add_h1("五、推广与前景类问答")

add_q('这个方案从“试点”到“全面推广”，最大的障碍是什么？')
add_a(
    "主要障碍不是技术演示，而是三件事："
    "① 制度与模板标准化——不同项目表格、流程、权限差异大；"
    "② 组织习惯——现场是否愿意从“等系统”转为“在平台上配流程”；"
    "③ 生产级非功能需求——权限分级、稳定运维、模型合规、与集团统一身份和数据标准对接。"
    "应对策略：先选 1–2 个成熟地盘做闭环试点，沉淀 SOP 和 Skill 包，再复制；"
    "IT 负责底座与安全，内派同事负责场景落地培训。"
)

add_q("有没有梳理过公司内还有哪些类似场景可以直接复用？")
add_a(
    "已梳理多类可复用场景："
    "① 视觉巡检与复查闭环；② 群消息隐患 intake 与台账；③ 每日/每周安全简报与外部风险提醒；"
    "④ 制度问答与表格智能预填；⑤ 整改通知/警告信 DOCX 生成；"
    "⑥ 飞书机器人确认卡审批；⑦ WhatsApp 本地归档查询与 SQL 分析。"
    "这些均可通过同一中台 + 不同 Skill/Workflow 组合实现，无需重复开发底层能力。"
)

add_q("如果要把这个方案产品化对外输出，还需要补齐哪些能力？")
add_a(
    "需补齐：多租户与项目级权限体系、统一身份认证对接、生产监控与 SLA、"
    "模型与数据合规认证、Workflow 可视化编排器、更完整的飞书/中海通企业集成、"
    "视觉误报反馈与策略自优化、实施交付标准包（部署脚本、培训手册、Skill 市场）。"
    "当前版本适合集团内试点与大赛展示，对外产品化需增加运维与商业化配套。"
)

add_h1("六、评分维度可能追问（风险 / 创新 / 是否为了 AI 而 AI）")

add_q("如何控制 AI 幻觉、数据隐私和算法偏见等风险？")
add_a(
    "四道防线："
    "① 输出默认是草稿/候选/卡片，不直接生效；"
    "② 高风险动作统一 Pending Confirmation，人工批准后才执行；"
    "③ LLM Gateway 脱敏 + 规则路由优先，减少模型胡判；"
    "④ 全链路 Audit Log / Job Store / Workflow Run 可追溯。"
    "视觉和分类结果保留证据图片与置信度，支持人工驳回和误报反馈。"
)

add_q('技术先进性与创新性体现在哪里？会不会被理解为“为了 AI 而 AI”？')
add_a(
    "创新点在于管理范式与工程化路径，而非堆概念："
    "① “数字化倒过来”——IT 建中台，现场共创流程；"
    "② Agentic Orchestration——Intent + Skill + Workflow + Tool Calling 统一编排；"
    "③ 多模态安全闭环——视觉、群聊、文档、制度、通知在同一平台协同；"
    "④ 可控 AI——Human-in-the-loop 写进架构而非口号。"
    "我们刻意避免“聊天框替代业务系统”，每个 AI 能力都对应明确痛点和人工确认点，因此不是为 AI 而 AI。"
)

add_q("汇报时如何在有限时间内讲清核心价值？")
add_a(
    "建议 8–10 分钟结构："
    "1 分钟讲痛点与范式反转；2 分钟讲三层架构一张图；"
    "3 分钟演示一条闭环（如视觉候选 → 制度检索 → 整改通知 Diff → 确认卡）；"
    "2 分钟讲可复制推广与价值；1 分钟回应风险与组织变革。"
    "口播金句：让最懂现场的人定义问题，让最懂业务的人驱动优化；"
    "赤瞳不替人做决定，而是把 AI 能力交到一线。"
)

add_h1("七、人文中海赛道相关问答（若适用）")

add_q('创作灵感来源是什么？AI 是“工具”还是“合作者”？')
add_a(
    "灵感来自一线安全管理的真实痛点：不是缺数据，而是缺把数据变成行动的协同机制。"
    "AI 在我们的实践中是“合作者”——负责理解、整理、生成草稿和提醒；"
    "人是“决策者”——负责确认、审批和责任归属。"
    "这与人文中海强调的“科技服务人、增强而非替代人”一致。"
)

add_q("为什么选择这款 AI 工具？有没有对比过其他工具？")
add_a(
    "不是单一工具，而是“OpenAI 兼容 LLM + 本地 VLM/OCR + RAG 向量库 + 自研中台编排”。"
    "对比纯 ChatGPT/纯 Copilot 类工具，赤瞳多了现场工具链、确认治理和工地场景 Workflow；"
    "对比传统安全软件，多了自然语言与多模态理解。"
    "选择依据是可落地、可审计、可扩展，而非追单一模型榜单。"
)

add_q("其他同事复用工作流需要多久？有没有 SOP？")
add_a(
    "基础使用（问答、简报、制度检索、文档 Diff 预览）可在 1–2 小时培训内上手；"
    "配置一个简单 Skill/Workflow 约 0.5–1 天；完整闭环落地需内派同事带教 1–2 周。"
    "我们正在沉淀：部署手册、Skill 说明、确认策略清单、试点 SOP，"
    "目标让“会描述场景的人”也能参与数字化。"
)

add_q('创作过程中，有没有“AI 不懂我”的瞬间？你是怎么解决的？')
add_a(
    "有，典型在繁体/英文混排制度、工地俚语、图片误判上。"
    "解决方式：① 先用 RAG 喂项目制度与历史案例；② 视觉结果只作候选，人工确认入库；"
    "③ 规则路由兜底关键 intent；④ 持续收集误报反馈优化提示词和策略包。"
    "我们接受 AI 会出错，所以架构上必须让人能轻松纠正。"
)

add_q("这个作品最让你自豪的细节是什么？最遗憾的缺陷是什么？")
add_a(
    "自豪：Human-in-the-loop 不是 PPT 概念，而是写进了 confirmation_service、卡片回传和审计链路里——"
    "AI 再强也不能绕过确认。"
    "遗憾：端到端“发现→通知→复查→关闭”完整状态机尚未在所有地盘生产验证，"
    "Workflow 数量和稳定性仍需继续打磨。"
    "我们如实呈现 MVP 阶段，但展示清晰方向与可复制路径。"
)

add_h1("八、建议口播收尾（30 秒）")
add_a(
    "传统数字化是“地盘提需求、IT 做系统”；AI 时代我们把流程倒过来："
    "IT 和中台团队搭建统一底座，让最懂现场的人定义问题，让最懂业务的人驱动优化。"
    "赤瞳不是替代管理决策，而是用人机协同把 AI 能力交到一线，实现平台赋能、现场共创。谢谢各位评委。"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("文档生成日期：2026-06-23  |  赤瞳安全智能平台项目组")
set_run_font(r, size=9, color=(120, 120, 120))

doc.save(str(out_path))
print(f"saved: {out_path}")
print(f"size: {out_path.stat().st_size}")
