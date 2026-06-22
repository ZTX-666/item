from __future__ import annotations

import importlib
from pathlib import Path

from docx import Document


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def _write_split_run_docx(path: Path) -> None:
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("202")
    para.add_run("4")
    para.add_run("年度安全巡检")
    doc.save(path)


def test_config_defaults_use_repo_relative_paths(monkeypatch):
    for name in [
        "AGENT_TOOLBOX_ROOT",
        "AGENT_TOOLBOX_WORKSPACE",
        "VLM_DETECTION_DIR",
        "RTMP_SNAPSHOT_SCRIPT",
        "REPORT_SCRIPT",
    ]:
        monkeypatch.setenv(name, "")

    import agent_toolbox.config as config

    config = importlib.reload(config)
    repo = config.ROOT.parent
    settings = config.Settings()

    assert settings.root == config.ROOT
    assert settings.workspace == config.ROOT / "workspace"
    assert settings.vlm_detection_dir == repo / "vlm-detection"
    assert settings.rtmp_snapshot_script == repo / "rtmp-tools" / "rtmp_snapshot.py"
    assert settings.report_script == repo / "report-generators" / "generate_community_doc.py"


def test_env_example_documents_optional_paths_and_preserves_publish3_wacli_notes():
    env_example = Path(__file__).resolve().parents[1] / ".env.example"
    text = env_example.read_text(encoding="utf-8")

    assert "Copy to .env ONLY if you need to override defaults." in text
    assert "automatically resolves them relative to the repository layout" in text
    assert "# VLM_DETECTION_DIR=" in text
    assert "# RTMP_SNAPSHOT_SCRIPT=" in text
    assert "# REPORT_SCRIPT=" in text
    assert "publish3.0 与 agent-toolbox 同在项目根目录" in text
    assert "# WACLI_BIN=E:\\publish3.0\\runtime\\bin\\wacli.exe" in text
    assert "# WACLI_WORKDIR=E:\\publish3.0\\runtime\\data" in text
    assert "AGENT_TOOLBOX_ROOT=J:" not in text


def test_apply_changeset_replaces_text_across_word_runs(tmp_path):
    from agent_toolbox.tools import docmate_docx

    docmate_docx._doc_store.clear()
    docmate_docx._changeset_store.clear()
    source = tmp_path / "split-runs.docx"
    output = tmp_path / "split-runs-updated.docx"
    _write_split_run_docx(source)

    read_result = docmate_docx.docmate_read_docx(
        docmate_docx.DocmateReadDocxRequest(file_path=str(source))
    )
    doc_id = read_result.data["doc_id"]
    changeset = docmate_docx.docmate_generate_changeset(
        docmate_docx.DocmateGenerateChangesetRequest(
            doc_id=doc_id,
            instruction="把2024改为2026",
        )
    )
    change_id = changeset.data["changes"][0]["change_id"]

    apply_result = docmate_docx.docmate_apply_changeset(
        docmate_docx.DocmateApplyChangesetRequest(
            changeset_id=changeset.data["changeset_id"],
            accepted_change_ids=[change_id],
            save_as=str(output),
        )
    )

    assert apply_result.ok is True
    assert apply_result.data["applied"] == 1
    assert apply_result.data["errors"] == []
    assert Document(output).paragraphs[0].text == "2026年度安全巡检"


def test_docmate_get_document_returns_loaded_paragraphs_and_stats(tmp_path):
    from agent_toolbox.tools import docmate_docx

    docmate_docx._doc_store.clear()
    docmate_docx._changeset_store.clear()
    source = tmp_path / "source.docx"
    _write_docx(source, ["第一段", "第二段"])

    read_result = docmate_docx.docmate_read_docx(
        docmate_docx.DocmateReadDocxRequest(file_path=str(source))
    )
    doc_id = read_result.data["doc_id"]

    assert hasattr(docmate_docx, "docmate_get_document")
    result = docmate_docx.docmate_get_document(doc_id)

    assert result.ok is True
    assert result.data["doc_id"] == doc_id
    assert result.data["paragraphs"] == [
        {"index": 0, "text": "第一段", "style": "Normal"},
        {"index": 1, "text": "第二段", "style": "Normal"},
    ]
    assert result.data["stats"]["paragraph_count"] == 2


def test_docmate_register_changeset_stores_llm_edits_for_preview_and_apply(tmp_path):
    from agent_toolbox.tools import docmate_docx

    docmate_docx._doc_store.clear()
    docmate_docx._changeset_store.clear()
    source = tmp_path / "register.docx"
    output = tmp_path / "register-updated.docx"
    _write_docx(source, ["旧内容需要复核"])

    read_result = docmate_docx.docmate_read_docx(
        docmate_docx.DocmateReadDocxRequest(file_path=str(source))
    )
    doc_id = read_result.data["doc_id"]

    assert hasattr(docmate_docx, "docmate_register_changeset")
    register_result = docmate_docx.docmate_register_changeset(
        doc_id,
        "AI 编辑建议",
        [
            {"type": "replace", "target": "旧内容", "replacement": "新内容"},
            {"type": "append", "replacement": "补充说明"},
        ],
    )

    assert register_result.ok is True
    assert register_result.data["doc_id"] == doc_id
    assert register_result.data["total_changes"] == 2

    preview_result = docmate_docx.docmate_preview_changeset(
        docmate_docx.DocmatePreviewChangesetRequest(
            changeset_id=register_result.data["changeset_id"]
        )
    )
    assert [card["type"] for card in preview_result.data["preview_cards"]] == [
        "text_replace",
        "text_append",
    ]

    apply_result = docmate_docx.docmate_apply_changeset(
        docmate_docx.DocmateApplyChangesetRequest(
            changeset_id=register_result.data["changeset_id"],
            accepted_change_ids=[
                change["change_id"] for change in register_result.data["changes"]
            ],
            save_as=str(output),
        )
    )
    updated = Document(output)

    assert apply_result.ok is True
    assert apply_result.data["applied"] == 2
    assert [para.text for para in updated.paragraphs] == ["新内容需要复核", "补充说明"]


def test_registry_and_http_routes_expose_docmate_document_tools():
    from agent_toolbox.app import app
    from agent_toolbox.registry import tool_specs

    specs = {spec.name: spec for spec in tool_specs()}
    paths = {route.path for route in app.routes}

    assert "docmate_get_document" in specs
    assert "docmate_register_changeset" in specs
    assert "/tools/docmate_get_document" in paths
    assert "/tools/docmate_register_changeset" in paths
