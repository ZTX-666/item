"""
docx_parser.py — 解析 .docx 结构：段落/表格/图片 → DocumentStructure

按《DocMate EHS docx升级与Skill落地实施方案 v1》第 9.1 节伪代码实现。
"""

from __future__ import annotations

import hashlib
import io
import os
import zipfile
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from .docx_model import (
    DocumentStructure,
    ParagraphModel,
    RunModel,
    TableModel,
    TableCellModel,
    ImageModel,
    make_hash_id,
)

# ─── 1. 段落解析 ────────────────────────────────────────────────

def _extract_runs(paragraph) -> list[RunModel]:
    """从 python-docx paragraph 提取 run 信息（格式 + 文字）。"""
    runs = []
    for i, run in enumerate(paragraph.runs):
        text = run.text
        if not text:
            continue
        run_id = f"R{i + 1}"
        font = run.font
        runs.append(RunModel(
            run_id=run_id,
            text=text,
            bold=font.bold or False,
            italic=font.italic or False,
            underline=font.underline or False,
            font_name=font.name or "",
            font_size=int(font.size.pt) if font.size else 12,
            color=str(font.color.rgb) if font.color and font.color.rgb else "",
        ))
    return runs


def parse_paragraphs(doc: DocxDocument) -> list[ParagraphModel]:
    """解析全部段落，生成 hash 锚点。"""
    paragraphs = []
    for idx, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        # 跳过完全空段落（但保留仅含空格的）
        # 生成 hash 锚点
        pid = make_hash_id("P", idx, text if text else f"empty_{idx}")
        style_name = para.style.name if para.style else "Normal"
        runs = _extract_runs(para)
        paragraphs.append(ParagraphModel(
            paragraph_id=pid,
            index=idx,
            text=text,
            style=style_name,
            runs=runs,
        ))
    return paragraphs


# ─── 2. 表格解析 ────────────────────────────────────────────────

def parse_tables(doc: DocxDocument) -> list[TableModel]:
    """解析全部表格，提取每个单元格的文字。"""
    tables = []
    for idx, table in enumerate(doc.tables, start=1):
        rows = len(table.rows)
        cols = len(table.columns)
        cells = []

        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(TableCellModel(
                        row=r_idx,
                        col=c_idx,
                        text=cell_text,
                    ))

        # 生成 table_id
        sample_text = cells[0].text if cells else f"table_{idx}"
        tid = make_hash_id("T", idx, sample_text)

        tables.append(TableModel(
            table_id=tid,
            index=idx,
            rows=rows,
            cols=cols,
            cells=cells,
        ))
    return tables


# ─── 3. 图片解析 ───────────────────────────────────────────────

def _extract_images_from_xml(doc: DocxDocument) -> list[dict]:
    """
    从 document.xml 中提取图片引用关系。
    返回 [{rId, anchor_paragraph_index, filename}]
    """
    images_info = []
    try:
        body = doc.element.body
        para_elements = body.findall(qn("w:p"))

        for p_idx, para in enumerate(para_elements, start=1):
            # 查找段落中的图片
            runs = para.findall(qn("w:r"))
            for run in runs:
                drawings = run.findall(qn("w:drawing"))
                for drawing in drawings:
                    # 深度搜索 blip 元素
                    blips = drawing.findall(".//" + qn("a:blip"))
                    for blip in blips:
                        embed = blip.get(qn("r:embed"))
                        if embed:
                            images_info.append({
                                "rId": embed,
                                "anchor_paragraph_index": p_idx,
                            })
    except Exception:
        pass

    return images_info


def parse_images(doc: DocxDocument, doc_source_path: str = "") -> list[ImageModel]:
    """
    解析文档中的图片：从 rels 提取文件名，结合 XML 锚点定位。
    """
    images = []
    try:
        # 1. 从 XML 获取图片引用
        xml_refs = _extract_images_from_xml(doc)

        # 2. 从 rels 获取图片关系
        part = doc.part
        rels = part.rels if hasattr(part, "rels") else {}

        for idx, ref in enumerate(xml_refs, start=1):
            rId = ref.get("rId", "")
            anchor_idx = ref.get("anchor_paragraph_index", 0)

            # 查找关系
            rel = rels.get(rId)
            if not rel:
                continue

            target_ref = str(rel.target_ref) if hasattr(rel, "target_ref") else ""
            filename = os.path.basename(target_ref)

            # 生成 image_id
            img_id = f"IMG{idx}"

            # 找到锚点段落的 paragraph_id
            anchor_pid = ""
            for p in parse_paragraphs(doc):
                if p.index == anchor_idx:
                    anchor_pid = p.paragraph_id
                    break

            images.append(ImageModel(
                image_id=img_id,
                index=idx,
                anchor_paragraph_id=anchor_pid,
                name=filename,
                width_px=640,
                height_px=360,
                image_path="",
            ))
    except Exception as e:
        # 图片解析失败不影响整体流程
        pass

    return images


# ─── 4. 主解析入口 ─────────────────────────────────────────────

def parse_docx(file_path: str) -> DocumentStructure:
    """
    主入口：解析 .docx 文件 → DocumentStructure。
    对应方案第 9.1 节伪代码。
    """
    file_path = str(Path(file_path).resolve())
    doc = DocxDocument(file_path)

    doc_id = hashlib.md5(file_path.encode()).hexdigest()[:12]

    paragraphs = parse_paragraphs(doc)
    tables = parse_tables(doc)
    images = parse_images(doc, file_path)

    return DocumentStructure(
        doc_id=doc_id,
        source_path=file_path,
        paragraphs=paragraphs,
        tables=tables,
        images=images,
    )


# ─── 5. 索引查找辅助 ───────────────────────────────────────────

def find_paragraph_by_text(structure: DocumentStructure, text_fragment: str) -> Optional[ParagraphModel]:
    """根据文本片段模糊匹配段落。"""
    for p in structure.paragraphs:
        if text_fragment in p.text:
            return p
    return None


def find_paragraph_by_style(structure: DocumentStructure, style_name: str) -> list[ParagraphModel]:
    """根据样式名查找所有匹配段落（如所有 Heading1）。"""
    return [p for p in structure.paragraphs if p.style == style_name]


def find_table_by_content(structure: DocumentStructure, keyword: str) -> Optional[TableModel]:
    """根据关键词查找包含该词的表格。"""
    for t in structure.tables:
        for cell in t.cells:
            if keyword in cell.text:
                return t
    return None
