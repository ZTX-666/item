from __future__ import annotations

import html
import json
import mimetypes
import os
import re
import subprocess
import urllib.parse
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

import requests
from pydantic import BaseModel, Field

from ..config import settings
from ..runner import CommandError, run_command

MAX_OUTPUT_CHARS = 32_000
MAX_FILE_BYTES = 8 * 1024 * 1024
MAX_URL_BYTES = 512 * 1024
TEXT_EXTENSIONS = {
    ".txt", ".md", ".json", ".jsonl", ".csv", ".tsv", ".yaml", ".yml", ".xml",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".vue", ".css", ".html", ".htm",
    ".bat", ".ps1", ".sh", ".sql", ".ini", ".env", ".log", ".toml",
}

BLOCKED_SHELL_PATTERNS = [
    r"\brm\s+-rf\b",
    r"\bdel\s+/[fq]",
    r"\bformat\s+[a-z]:",
    r"\bshutdown\b",
    r"\brestart-computer\b",
    r"\bstop-computer\b",
    r"\breg\s+delete\b",
    r"\bRemove-Item\s+.*-Recurse\s+-Force",
    r"\bInvoke-Expression\b",
    r"\biex\b",
    r"\bwget\s+.*\|\s*sh",
    r"\bcurl\s+.*\|\s*sh",
    r"\|\s*powershell",
    r">\s*/dev/",
    r"\bmkfs\b",
    r"\bdd\s+if=",
]

USER_AGENT = "ChitungDesktopAutomation/1.0 (+https://chitung.local; safety-platform)"


class WebSearchRequest(BaseModel):
    query: str = Field(..., min_length=1, max_length=500)
    limit: int = Field(default=8, ge=1, le=15)
    region: str = Field(default="wt-wt", description="DuckDuckGo region code, e.g. cn-zh for Chinese results")


class FetchUrlRequest(BaseModel):
    url: str = Field(..., min_length=4, max_length=2048)
    extract_mode: str = Field(default="readable", description="readable | links | raw")
    max_chars: int = Field(default=12000, ge=500, le=50000)


class RunBashRequest(BaseModel):
    command: str = Field(..., min_length=1, max_length=4000)
    cwd: str | None = None
    timeout_seconds: int = Field(default=120, ge=5, le=600)


class RunPowerShellRequest(BaseModel):
    command: str = Field(..., min_length=1, max_length=4000)
    cwd: str | None = None
    timeout_seconds: int = Field(default=120, ge=5, le=600)


class ReadLocalFileRequest(BaseModel):
    path: str = Field(..., min_length=1, max_length=2048)
    max_chars: int = Field(default=16000, ge=500, le=80000)
    include_binary_preview: bool = False


class _DdgResultParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.results: list[dict[str, str]] = []
        self._capture: str | None = None
        self._current: dict[str, str] = {}
        self._text_parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_map = {key: (value or "") for key, value in attrs}
        if tag == "a" and "result__a" in (attrs_map.get("class") or ""):
            self._capture = "title"
            self._current = {"url": attrs_map.get("href", ""), "title": "", "snippet": ""}
            self._text_parts = []
        elif tag == "a" and "result__snippet" in (attrs_map.get("class") or ""):
            self._capture = "snippet"
            self._text_parts = []

    def handle_data(self, data: str) -> None:
        if self._capture:
            self._text_parts.append(data)

    def handle_endtag(self, tag: str) -> None:
        if tag != "a" or not self._capture:
            return
        text = html.unescape("".join(self._text_parts).strip())
        if self._capture == "title":
            self._current["title"] = text
        elif self._capture == "snippet":
            self._current["snippet"] = text
            if self._current.get("title"):
                self.results.append(self._current)
            self._current = {}
        self._capture = None
        self._text_parts = []


def _allowed_read_roots() -> list[Path]:
    roots: list[Path] = []
    for candidate in (settings.workspace, settings.root, settings.root.parent):
        try:
            resolved = candidate.expanduser().resolve()
        except OSError:
            continue
        if resolved.exists() and resolved not in roots:
            roots.append(resolved)
    for part in os.getenv("DESKTOP_ALLOWED_READ_ROOTS", "").split(";"):
        part = part.strip()
        if not part:
            continue
        try:
            resolved = Path(part).expanduser().resolve()
        except OSError:
            continue
        if resolved.exists() and resolved not in roots:
            roots.append(resolved)
    return roots


def _resolve_allowed_path(raw_path: str, *, write_ok: bool = False) -> Path:
    candidate = Path(raw_path.strip().strip('"')).expanduser()
    if not candidate.is_absolute():
        candidate = (settings.workspace / candidate).resolve()
    else:
        candidate = candidate.resolve()
    allowed = False
    for root in _allowed_read_roots():
        try:
            candidate.relative_to(root)
            allowed = True
            break
        except ValueError:
            continue
    if not allowed:
        raise ValueError(f"路径不在允许范围内：{candidate}")
    if not write_ok and not candidate.exists():
        raise FileNotFoundError(f"文件不存在：{candidate}")
    sensitive_names = {".env", "credentials.json", "secrets.json", "id_rsa", "id_dsa"}
    if candidate.name.lower() in sensitive_names or candidate.suffix.lower() in {".pem", ".key", ".pfx"}:
        raise ValueError(f"拒绝读取敏感文件：{candidate.name}")
    return candidate


def _shell_workspace(cwd: str | None) -> Path:
    if cwd:
        return _resolve_allowed_path(cwd, write_ok=True)
    settings.workspace.mkdir(parents=True, exist_ok=True)
    return settings.workspace.resolve()


def _validate_shell_command(command: str) -> None:
    lowered = command.lower()
    for pattern in BLOCKED_SHELL_PATTERNS:
        if re.search(pattern, lowered, flags=re.IGNORECASE):
            raise ValueError(f"命令包含受限操作，已拒绝：{pattern}")
    if len(command.strip()) > 4000:
        raise ValueError("命令过长")


def _truncate(text: str, limit: int = MAX_OUTPUT_CHARS) -> tuple[str, bool]:
    if len(text) <= limit:
        return text, False
    return text[: limit - 20] + "\n...[truncated]", True


def _strip_html(raw: str) -> str:
    raw = re.sub(r"(?is)<(script|style|noscript).*?>.*?</\1>", " ", raw)
    raw = re.sub(r"(?is)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?is)</p\s*>", "\n\n", raw)
    raw = re.sub(r"(?is)<[^>]+>", " ", raw)
    raw = html.unescape(re.sub(r"[ \t\f\v]+", " ", raw))
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()


def web_search(payload: WebSearchRequest) -> dict[str, Any]:
    if not settings.desktop_automation_enabled:
        return {"ok": False, "error": "桌面自动化能力未启用（DESKTOP_AUTOMATION_ENABLED=false）"}
    query = payload.query.strip()
    data = {"q": query, "kl": payload.region}
    headers = {"User-Agent": USER_AGENT}
    try:
        response = requests.post(
            "https://html.duckduckgo.com/html/",
            data=data,
            headers=headers,
            timeout=settings.desktop_http_timeout_seconds,
        )
        response.raise_for_status()
    except requests.RequestException as exc:
        return {"ok": False, "error": f"网页搜索请求失败：{exc}", "query": query}

    parser = _DdgResultParser()
    parser.feed(response.text)
    results = parser.results[: payload.limit]
    return {
        "ok": True,
        "tool": "web_search",
        "query": query,
        "provider": "duckduckgo_html",
        "result_count": len(results),
        "results": results,
    }


def fetch_url_content(payload: FetchUrlRequest) -> dict[str, Any]:
    if not settings.desktop_automation_enabled:
        return {"ok": False, "error": "桌面自动化能力未启用（DESKTOP_AUTOMATION_ENABLED=false）"}
    parsed = urllib.parse.urlparse(payload.url.strip())
    if parsed.scheme not in {"http", "https"}:
        return {"ok": False, "error": "仅支持 http/https URL"}
    hostname = (parsed.hostname or "").lower()
    blocked_hosts = {"127.0.0.1", "localhost", "0.0.0.0", "::1"}
    if hostname in blocked_hosts or hostname.endswith(".local"):
        return {"ok": False, "error": "不允许抓取本地/内网地址"}

    headers = {"User-Agent": USER_AGENT}
    try:
        response = requests.get(
            payload.url,
            headers=headers,
            timeout=settings.desktop_http_timeout_seconds,
            allow_redirects=True,
        )
        response.raise_for_status()
    except requests.RequestException as exc:
        return {"ok": False, "error": f"URL 抓取失败：{exc}", "url": payload.url}

    content_type = response.headers.get("content-type", "text/html").split(";")[0].strip()
    raw_bytes = response.content[:MAX_URL_BYTES]
    truncated_bytes = len(response.content) > MAX_URL_BYTES

    if payload.extract_mode == "raw":
        text = raw_bytes.decode(response.encoding or "utf-8", errors="replace")
        content, was_truncated = _truncate(text, payload.max_chars)
        return {
            "ok": True,
            "tool": "fetch_url_content",
            "url": payload.url,
            "content_type": content_type,
            "mode": "raw",
            "truncated": truncated_bytes or was_truncated,
            "content": content,
        }

    if "html" in content_type:
        readable = _strip_html(raw_bytes.decode(response.encoding or "utf-8", errors="replace"))
        links = re.findall(r'href=["\\\'](https?://[^"\\\']+)["\\\']', response.text)
        links = list(dict.fromkeys(links))[:20]
        if payload.extract_mode == "links":
            return {
                "ok": True,
                "tool": "fetch_url_content",
                "url": payload.url,
                "content_type": content_type,
                "mode": "links",
                "link_count": len(links),
                "links": links,
            }
        content, was_truncated = _truncate(readable, payload.max_chars)
        return {
            "ok": True,
            "tool": "fetch_url_content",
            "url": payload.url,
            "content_type": content_type,
            "mode": "readable",
            "title": _extract_title(response.text),
            "truncated": truncated_bytes or was_truncated,
            "content": content,
            "links": links[:8],
        }

    if content_type.startswith("application/json"):
        try:
            parsed_json = json.loads(raw_bytes.decode("utf-8", errors="replace"))
            text = json.dumps(parsed_json, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            text = raw_bytes.decode("utf-8", errors="replace")
    else:
        text = raw_bytes.decode(response.encoding or "utf-8", errors="replace")

    content, was_truncated = _truncate(text, payload.max_chars)
    return {
        "ok": True,
        "tool": "fetch_url_content",
        "url": payload.url,
        "content_type": content_type,
        "mode": payload.extract_mode,
        "truncated": truncated_bytes or was_truncated,
        "content": content,
    }


def _extract_title(raw_html: str) -> str:
    match = re.search(r"(?is)<title[^>]*>(.*?)</title>", raw_html)
    if not match:
        return ""
    return html.unescape(re.sub(r"\s+", " ", match.group(1)).strip())


def run_bash_command(payload: RunBashRequest) -> dict[str, Any]:
    if not settings.desktop_automation_enabled:
        return {"ok": False, "error": "桌面自动化能力未启用（DESKTOP_AUTOMATION_ENABLED=false）"}
    _validate_shell_command(payload.command)
    cwd = _shell_workspace(payload.cwd)
    if os.name == "nt":
        shell_cmd = ["bash", "-lc", payload.command]
    else:
        shell_cmd = ["bash", "-lc", payload.command]
    return _run_shell("run_bash_command", shell_cmd, cwd, payload.timeout_seconds)


def run_powershell_command(payload: RunPowerShellRequest) -> dict[str, Any]:
    if not settings.desktop_automation_enabled:
        return {"ok": False, "error": "桌面自动化能力未启用（DESKTOP_AUTOMATION_ENABLED=false）"}
    _validate_shell_command(payload.command)
    cwd = _shell_workspace(payload.cwd)
    shell_cmd = [
        "powershell",
        "-NoProfile",
        "-NonInteractive",
        "-ExecutionPolicy",
        "Bypass",
        "-Command",
        payload.command,
    ]
    return _run_shell("run_powershell_command", shell_cmd, cwd, payload.timeout_seconds)


def _run_shell(tool: str, command: list[str], cwd: Path, timeout_seconds: int) -> dict[str, Any]:
    try:
        result = run_command(command, cwd=cwd, timeout=timeout_seconds)
        stdout, truncated = _truncate(result.stdout or "")
        stderr, stderr_truncated = _truncate(result.stderr or "")
        return {
            "ok": True,
            "tool": tool,
            "command": command,
            "cwd": str(cwd),
            "exit_code": result.returncode,
            "stdout": stdout,
            "stderr": stderr,
            "truncated": truncated or stderr_truncated,
        }
    except CommandError as exc:
        stdout, truncated = _truncate(exc.stdout or "")
        stderr, stderr_truncated = _truncate(exc.stderr or "")
        return {
            "ok": False,
            "tool": tool,
            "command": exc.command,
            "cwd": str(cwd),
            "exit_code": exc.returncode,
            "stdout": stdout,
            "stderr": stderr,
            "truncated": truncated or stderr_truncated,
            "error": str(exc),
        }
    except subprocess.TimeoutExpired:
        return {"ok": False, "tool": tool, "command": command, "cwd": str(cwd), "error": "命令执行超时"}
    except FileNotFoundError:
        return {"ok": False, "tool": tool, "error": "未找到 bash 或 PowerShell 可执行文件，请确认系统环境。"}
    except ValueError as exc:
        return {"ok": False, "tool": tool, "error": str(exc)}


def read_local_file(payload: ReadLocalFileRequest) -> dict[str, Any]:
    if not settings.desktop_automation_enabled:
        return {"ok": False, "error": "桌面自动化能力未启用（DESKTOP_AUTOMATION_ENABLED=false）"}
    try:
        path = _resolve_allowed_path(payload.path)
    except (ValueError, FileNotFoundError) as exc:
        return {"ok": False, "tool": "read_local_file", "error": str(exc)}

    if path.is_dir():
        entries = sorted(path.iterdir(), key=lambda item: item.name.lower())[:200]
        return {
            "ok": True,
            "tool": "read_local_file",
            "path": str(path),
            "kind": "directory",
            "entry_count": len(entries),
            "entries": [{"name": item.name, "is_dir": item.is_dir()} for item in entries],
        }

    size = path.stat().st_size
    if size > MAX_FILE_BYTES:
        return {
            "ok": False,
            "tool": "read_local_file",
            "path": str(path),
            "error": f"文件过大（{size} bytes），上限 {MAX_FILE_BYTES} bytes",
        }

    suffix = path.suffix.lower()
    mime, _ = mimetypes.guess_type(str(path))

    if suffix == ".ipynb" or mime == "application/x-ipynb+json":
        return _read_ipynb(path, payload.max_chars)
    if suffix == ".docx":
        return _read_docx(path, payload.max_chars)
    if suffix == ".pdf" or mime == "application/pdf":
        return _read_pdf(path, payload.max_chars)
    if suffix in TEXT_EXTENSIONS or (mime and mime.startswith("text/")):
        text = path.read_text(encoding="utf-8", errors="replace")
        content, truncated = _truncate(text, payload.max_chars)
        return {
            "ok": True,
            "tool": "read_local_file",
            "path": str(path),
            "kind": "text",
            "mime": mime,
            "size": size,
            "truncated": truncated,
            "content": content,
        }
    if mime and mime.startswith("image/"):
        return {
            "ok": True,
            "tool": "read_local_file",
            "path": str(path),
            "kind": "image",
            "mime": mime,
            "size": size,
            "note": "图片文件已识别；如需 OCR 请使用 ocr_document_or_image。",
        }

    if payload.include_binary_preview:
        preview = path.read_bytes()[:256]
        return {
            "ok": True,
            "tool": "read_local_file",
            "path": str(path),
            "kind": "binary",
            "mime": mime,
            "size": size,
            "preview_hex": preview.hex(),
        }
    return {
        "ok": False,
        "tool": "read_local_file",
        "path": str(path),
        "error": f"暂不支持直接读取该格式：{suffix or mime or 'unknown'}",
    }


def _read_ipynb(path: Path, max_chars: int) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    cells = data.get("cells") if isinstance(data, dict) else []
    parts: list[str] = []
    for index, cell in enumerate(cells or []):
        if not isinstance(cell, dict):
            continue
        source = cell.get("source")
        if isinstance(source, list):
            text = "".join(str(item) for item in source)
        else:
            text = str(source or "")
        parts.append(f"[{index}] {cell.get('cell_type', 'cell')}\n{text}".strip())
    content, truncated = _truncate("\n\n".join(parts), max_chars)
    return {
        "ok": True,
        "tool": "read_local_file",
        "path": str(path),
        "kind": "ipynb",
        "cell_count": len(cells or []),
        "truncated": truncated,
        "content": content,
    }


def _read_docx(path: Path, max_chars: int) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError:
        return {"ok": False, "tool": "read_local_file", "path": str(path), "error": "缺少 python-docx 依赖"}
    document = Document(str(path))
    paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    content, truncated = _truncate("\n".join(paragraphs), max_chars)
    return {
        "ok": True,
        "tool": "read_local_file",
        "path": str(path),
        "kind": "docx",
        "paragraph_count": len(paragraphs),
        "truncated": truncated,
        "content": content,
    }


def _read_pdf(path: Path, max_chars: int) -> dict[str, Any]:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return {"ok": False, "tool": "read_local_file", "path": str(path), "error": "缺少 pypdfium2 依赖"}
    pdf = pdfium.PdfDocument(str(path))
    parts: list[str] = []
    for index in range(len(pdf)):
        page = pdf[index]
        textpage = page.get_textpage()
        parts.append(textpage.get_text_range() or "")
        if sum(len(part) for part in parts) >= max_chars:
            break
    content, truncated = _truncate("\n\n".join(part.strip() for part in parts if part.strip()), max_chars)
    return {
        "ok": True,
        "tool": "read_local_file",
        "path": str(path),
        "kind": "pdf",
        "page_count": len(pdf),
        "truncated": truncated,
        "content": content,
    }


def desktop_automation_health() -> dict[str, Any]:
    roots = [str(item) for item in _allowed_read_roots()]
    bash_ok = False
    powershell_ok = False
    try:
        bash_ok = bool(run_command(["bash", "--version"], timeout=10).stdout)
    except (CommandError, FileNotFoundError, subprocess.TimeoutExpired, TypeError):
        bash_ok = False
    try:
        powershell_ok = bool(run_command(["powershell", "-NoProfile", "-Command", "$PSVersionTable.PSVersion"], timeout=10).stdout)
    except (CommandError, FileNotFoundError, subprocess.TimeoutExpired, TypeError):
        powershell_ok = False
    return {
        "available": settings.desktop_automation_enabled,
        "enabled": settings.desktop_automation_enabled,
        "workspace": str(settings.workspace),
        "allowed_read_roots": roots,
        "bash_available": bash_ok,
        "powershell_available": powershell_ok,
        "http_timeout_seconds": settings.desktop_http_timeout_seconds,
    }
