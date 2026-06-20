"""DocMate Pipeline — AI-powered DOCX read → generate → preview → apply.

Provides four pipeline tools that form the DocMate (闪闪文档) intelligent
document editing workflow:

  read_docx       → parse .docx structure (paragraphs, tables, images)
  generate_changeset → analyze instruction, propose LLM-style changes
  preview_changeset  → return change preview cards
  apply_changeset    → write accepted changes to a new .docx with backup

All state is held in-memory (_doc_store / _changeset_store).
"""

from __future__ import annotations

import copy
import json
import re
import shutil
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

from pydantic import BaseModel, Field

from ..models import ToolResult, ToolSpec
from ..tasks import new_task_id, record_task_event


# ---------------------------------------------------------------------------
# In-memory state stores
# ---------------------------------------------------------------------------

_doc_store: dict[str, dict[str, Any]] = {}
_changeset_store: dict[str, dict[str, Any]] = {}


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


# ---------------------------------------------------------------------------
# Request models
# ---------------------------------------------------------------------------

class DocmateReadDocxRequest(BaseModel):
    file_path: str = Field(..., description="Absolute path to .docx file")


class DocmateGenerateChangesetRequest(BaseModel):
    doc_id: str = Field(..., description="Document ID from read_docx")
    instruction: str = Field(..., min_length=1, description="Natural language edit instruction")
    context: str = Field(default="", description="Optional extra context for the LLM")


class DocmatePreviewChangesetRequest(BaseModel):
    changeset_id: str = Field(..., description="Changeset ID from generate_changeset")


class DocmateApplyChangesetRequest(BaseModel):
    changeset_id: str = Field(..., description="Changeset ID to apply")
    accepted_change_ids: list[str] = Field(default_factory=list, description="IDs of changes to accept")
    save_as: str = Field(default="", description="Output path; defaults to source_modified.docx")


# ---------------------------------------------------------------------------
# Document reading
# ---------------------------------------------------------------------------

def docmate_read_docx(req: DocmateReadDocxRequest) -> ToolResult:
    """Parse a .docx file and return its structural overview."""
    task_id = new_task_id("docmate")
    fp = Path(req.file_path)

    if not fp.exists():
        return ToolResult(
            ok=False, tool="docmate_read_docx", task_id=task_id,
            summary=f"文件不存在: {req.file_path}",
            error="file_not_found",
        )

    if fp.suffix.lower() != ".docx":
        return ToolResult(
            ok=False, tool="docmate_read_docx", task_id=task_id,
            summary=f"不支持的文件类型: {fp.suffix}，仅支持 .docx",
            error="unsupported_format",
        )

    try:
        from docx import Document
        doc = Document(str(fp))
    except Exception as exc:
        return ToolResult(
            ok=False, tool="docmate_read_docx", task_id=task_id,
            summary=f"无法打开文档: {exc}",
            error=str(exc),
        )

    doc_id = str(uuid.uuid4())[:12]

    # Extract paragraphs
    paragraphs: list[dict[str, Any]] = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append({
            "index": idx,
            "text": text,
            "style": str(para.style.name) if para.style else "Normal",
            "type": "paragraph",
        })

    # Extract tables
    tables: list[dict[str, Any]] = []
    for t_idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)
        tables.append({
            "index": t_idx,
            "rows": len(rows_data),
            "cols": len(rows_data[0]) if rows_data else 0,
            "data": rows_data[:20],  # limit preview
            "type": "table",
        })
        # Inject table refs into paragraph list at approximate position
        if paragraphs:
            insert_at = min(len(paragraphs), t_idx * 5)
            paragraphs.insert(insert_at, {"index": -1, "text": f"[表格 {t_idx + 1}]", "style": "", "type": "table_ref"})

    # Count images
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    image_count = len(doc.part.rels) if hasattr(doc, 'part') else 0
    try:
        image_count = sum(1 for rel in doc.part.rels.values() if "image" in str(rel.reltype))
    except Exception:
        image_count = 0

    stats = {
        "paragraph_count": len([p for p in paragraphs if p.get("type") != "table_ref"]),
        "table_count": len(tables),
        "image_count": image_count,
        "total_chars": sum(len(p["text"]) for p in paragraphs if p.get("type") != "table_ref"),
    }

    structure = {
        "paragraphs": paragraphs,
        "tables": tables,
    }

    data = {
        "doc_id": doc_id,
        "source_path": str(fp),
        "filename": fp.name,
        "stats": stats,
        "structure": structure,
    }

    # Store
    _doc_store[doc_id] = {
        "source_path": str(fp),
        "filename": fp.name,
        "structure": structure,
        "doc_object": doc,
        "stats": stats,
        "read_at": _now(),
    }

    record_task_event(task_id, {"tool": "docmate_read_docx", "file": fp.name, "status": "ok"})

    return ToolResult(
        ok=True, tool="docmate_read_docx", task_id=task_id,
        summary=f"已解析文档「{fp.name}」：{stats['paragraph_count']} 段落, {stats['table_count']} 表格, {image_count} 图片。",
        data=data,
    )


# ---------------------------------------------------------------------------
# Change generation (smart pattern matching, no LLM required)
# ---------------------------------------------------------------------------

_QUOTE_RE = r'["\u201c\u201d\u300c\u300d\u2018\u2019]?'

# Regex patterns for common Chinese edit instructions
# Group 1 = target text, Group 2 = replacement (greedy to capture full phrase)
_REPLACE_RE = re.compile(
    r'(?:把|将)\s*' + _QUOTE_RE + r'(.+?)' + _QUOTE_RE
    + r'\s*(?:改[为成]|替换[为成]|修改[为成])\s*'
    + _QUOTE_RE + r'(.+)' + _QUOTE_RE
)
_REPLACE2_RE = re.compile(
    r'(?:把|将)\s*' + _QUOTE_RE + r'(.+?)' + _QUOTE_RE
    + r'\s*替换[为成]\s*'
    + _QUOTE_RE + r'(.+)' + _QUOTE_RE
)
_REPLACE_INV_RE = re.compile(
    r'用\s*' + _QUOTE_RE + r'(.+?)' + _QUOTE_RE
    + r'\s*替换\s*'
    + _QUOTE_RE + r'(.+)' + _QUOTE_RE
)
_BARE_REPLACE_RE = re.compile(
    _QUOTE_RE + r'(.+?)' + _QUOTE_RE
    + r'\s*改[为成]\s*'
    + _QUOTE_RE + r'(.+)' + _QUOTE_RE
)
_DELETE_RE = re.compile(r'(?:删除|移除|去掉)\s*' + _QUOTE_RE + r'(.+?)' + _QUOTE_RE)
_INSERT_RE = re.compile(
    r'在\s*' + _QUOTE_RE + r'(.+?)' + _QUOTE_RE
    + r'\s*[之后前]\s*(?:添加|插入|增加)\s*'
    + _QUOTE_RE + r'(.+)' + _QUOTE_RE
)
_APPEND_RE = re.compile(r'(?:新增|添加|追加)\s*' + _QUOTE_RE + r'(.+)' + _QUOTE_RE)

# Clause splitters
_CLAUSE_SPLIT_RE = re.compile(r'[，。；！？\n,;!?]+')

_RISK_KEYWORDS_HIGH = frozenset([
    "删除", "移除", "高风险", "重大", "停工", "事故", "死亡",
    "安全", "危险", "紧急", "立即", "禁止",
])
_RISK_KEYWORDS_MEDIUM = frozenset([
    "修改", "变更", "调整", "更新", "替换", "改为",
])


def _match_instruction(instruction: str) -> list[dict[str, Any]]:
    """Parse instruction into one or more change proposals.

    Splits the instruction into clauses (by Chinese/English punctuation),
    then applies regex patterns to each clause independently.
    """
    proposals: list[dict[str, Any]] = []
    offset = 0

    for clause in _CLAUSE_SPLIT_RE.split(instruction):
        clause = clause.strip()
        if not clause:
            offset += len(clause) + 1
            continue

        # Try patterns in priority order
        m = _REPLACE_RE.search(clause) or _REPLACE2_RE.search(clause)
        if m:
            proposals.append({
                "type": "replace",
                "target": m.group(1).strip(),
                "replacement": m.group(2).strip(),
                "match_span": (offset + m.start(), offset + m.end()),
            })
            offset += len(clause) + 1
            continue

        m = _REPLACE_INV_RE.search(clause)
        if m:
            proposals.append({
                "type": "replace",
                "target": m.group(2).strip(),
                "replacement": m.group(1).strip(),
                "match_span": (offset + m.start(), offset + m.end()),
            })
            offset += len(clause) + 1
            continue

        m = _BARE_REPLACE_RE.search(clause)
        if m:
            proposals.append({
                "type": "replace",
                "target": m.group(1).strip(),
                "replacement": m.group(2).strip(),
                "match_span": (offset + m.start(), offset + m.end()),
            })
            offset += len(clause) + 1
            continue

        m = _DELETE_RE.search(clause)
        if m:
            proposals.append({
                "type": "delete",
                "target": m.group(1).strip(),
                "replacement": "",
                "match_span": (offset + m.start(), offset + m.end()),
            })
            offset += len(clause) + 1
            continue

        m = _INSERT_RE.search(clause)
        if m:
            target = m.group(1).strip()
            extra = m.group(2).strip()
            proposals.append({
                "type": "replace",
                "target": target,
                "replacement": f"{target}{extra}",
                "note": f"在「{target}」后添加「{extra}」",
                "match_span": (offset + m.start(), offset + m.end()),
            })
            offset += len(clause) + 1
            continue

        m = _APPEND_RE.search(clause)
        if m:
            proposals.append({
                "type": "append",
                "target": "",
                "replacement": m.group(1).strip(),
                "match_span": (offset + m.start(), offset + m.end()),
            })
            offset += len(clause) + 1
            continue

        offset += len(clause) + 1

    # Deduplicate by span
    seen_spans: set[tuple[int, int]] = set()
    unique: list[dict[str, Any]] = []
    for p in proposals:
        span = p.pop("match_span")
        if span not in seen_spans:
            seen_spans.add(span)
            unique.append(p)

    return unique


def _find_in_text(target: str, paragraphs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Find paragraphs containing the target text."""
    matches: list[dict[str, Any]] = []
    for para in paragraphs:
        if para.get("type") == "table_ref":
            continue
        text = para.get("text", "")
        if target and target in text:
            matches.append(para)
    return matches


def _assess_risk(instruction: str, target: str) -> str:
    """Assess risk level of a change."""
    combined = instruction + target
    for kw in _RISK_KEYWORDS_HIGH:
        if kw in combined:
            return "high"
    for kw in _RISK_KEYWORDS_MEDIUM:
        if kw in combined:
            return "medium"
    return "low"


def _make_changeset(proposals: list[dict[str, Any]], paragraphs: list[dict[str, Any]], instruction: str) -> dict[str, Any]:
    """Convert proposals into a changeset with preview cards."""
    changes: list[dict[str, Any]] = []
    preview_cards: list[dict[str, Any]] = []

    for i, prop in enumerate(proposals):
        change_id = str(uuid.uuid4())[:8]
        ptype = prop["type"]
        target = prop.get("target", "")
        replacement = prop.get("replacement", "")

        if ptype == "replace":
            matches = _find_in_text(target, paragraphs)
            if matches:
                # Use first match
                before_text = target
                after_text = replacement
                confidence = 0.85 + (0.1 if len(matches) == 1 else 0)
                risk = _assess_risk(instruction, target)

                change_entry = {
                    "change_id": change_id,
                    "type": "text_replace",
                    "target": target,
                    "replacement": replacement,
                    "paragraph_index": matches[0]["index"],
                    "context_before": matches[0]["text"][:80],
                    "occurrences": len(matches),
                }
                changes.append(change_entry)

                preview_cards.append({
                    "change_id": change_id,
                    "title": f"替换「{_truncate(target, 30)}」→「{_truncate(replacement, 30)}」",
                    "type": "text_replace",
                    "before": _truncate(matches[0]["text"], 80),
                    "after": matches[0]["text"].replace(target, replacement, 1)[:80],
                    "risk_level": risk,
                    "confidence": round(confidence, 2),
                    "occurrences": len(matches),
                    "paragraph_index": matches[0]["index"],
                })
            else:
                # Target not found in text
                changes.append({
                    "change_id": change_id,
                    "type": "text_replace",
                    "target": target,
                    "replacement": replacement,
                    "paragraph_index": -1,
                    "occurrences": 0,
                })
                preview_cards.append({
                    "change_id": change_id,
                    "title": f"替换「{_truncate(target, 30)}」→「{_truncate(replacement, 30)}」（未匹配）",
                    "type": "text_replace",
                    "before": f"(未在文档中找到「{_truncate(target, 40)}」)",
                    "after": f"将尝试全文档搜索替换",
                    "risk_level": "low",
                    "confidence": 0.45,
                    "occurrences": 0,
                    "paragraph_index": -1,
                })

        elif ptype == "delete":
            matches = _find_in_text(target, paragraphs)
            risk = "high"  # deletions are always high risk
            confidence = 0.8

            changes.append({
                "change_id": change_id,
                "type": "text_delete",
                "target": target,
                "replacement": "",
                "paragraph_index": matches[0]["index"] if matches else -1,
                "occurrences": len(matches),
            })

            preview_cards.append({
                "change_id": change_id,
                "title": f"删除「{_truncate(target, 30)}」",
                "type": "text_delete",
                "before": matches[0]["text"][:80] if matches else f"「{_truncate(target, 40)}」",
                "after": matches[0]["text"].replace(target, "", 1)[:80] if matches else "",
                "risk_level": risk,
                "confidence": round(confidence, 2),
                "occurrences": len(matches),
                "paragraph_index": matches[0]["index"] if matches else -1,
            })

        elif ptype == "append":
            risk = "low"
            confidence = 0.9

            changes.append({
                "change_id": change_id,
                "type": "text_append",
                "target": "",
                "replacement": replacement,
                "paragraph_index": -1,
                "occurrences": 0,
            })

            preview_cards.append({
                "change_id": change_id,
                "title": f"新增「{_truncate(replacement, 30)}」",
                "type": "text_append",
                "before": "(文档末尾)",
                "after": replacement,
                "risk_level": risk,
                "confidence": round(confidence, 2),
                "occurrences": 0,
                "paragraph_index": -1,
            })

    return {"changes": changes, "preview_cards": preview_cards}


def _truncate(s: str, n: int) -> str:
    if len(s) <= n:
        return s
    return s[:n-3] + "..."


def docmate_generate_changeset(req: DocmateGenerateChangesetRequest) -> ToolResult:
    """Analyze instruction against loaded document and produce change proposals."""
    task_id = new_task_id("docmate")

    doc = _doc_store.get(req.doc_id)
    if not doc:
        return ToolResult(
            ok=False, tool="docmate_generate_changeset", task_id=task_id,
            summary=f"文档不存在: {req.doc_id}。请先调用 docmate_read_docx。",
            error="doc_not_found",
        )

    paragraphs = [
        p for p in doc["structure"]["paragraphs"]
        if p.get("type") != "table_ref"
    ]

    # Parse instruction into proposals
    proposals = _match_instruction(req.instruction)

    if not proposals:
        # Fallback: treat the whole instruction as a fuzzy search
        # Try keyword-based match
        words = [w for w in re.findall(r"[\w\u4e00-\u9fff]+", req.instruction) if len(w) >= 2]
        if words:
            # Use first meaningful word as search target
            proposals = [{"type": "replace", "target": words[0], "replacement": req.instruction}]

    if not proposals:
        return ToolResult(
            ok=False, tool="docmate_generate_changeset", task_id=task_id,
            summary="无法解析编辑指令，请使用更明确的表述（如「把X改为Y」）。",
            error="unparseable_instruction",
        )

    # Build changeset
    changeset_data = _make_changeset(proposals, paragraphs, req.instruction)
    changeset_id = str(uuid.uuid4())[:12]

    changeset = {
        "changeset_id": changeset_id,
        "doc_id": req.doc_id,
        "instruction": req.instruction,
        "context": req.context,
        "changes": changeset_data["changes"],
        "preview_cards": changeset_data["preview_cards"],
        "generated_at": _now(),
    }

    _changeset_store[changeset_id] = changeset

    record_task_event(task_id, {
        "tool": "docmate_generate_changeset",
        "doc_id": req.doc_id,
        "changes": len(changeset_data["changes"]),
        "status": "ok",
    })

    return ToolResult(
        ok=True, tool="docmate_generate_changeset", task_id=task_id,
        summary=f"已生成 {len(changeset_data['preview_cards'])} 项修改方案。",
        data={
            "changeset_id": changeset_id,
            "doc_id": req.doc_id,
            "instruction": req.instruction,
            "changes": changeset_data["changes"],
            "preview_cards": changeset_data["preview_cards"],
            "total_changes": len(changeset_data["preview_cards"]),
        },
    )


# ---------------------------------------------------------------------------
# Preview
# ---------------------------------------------------------------------------

def docmate_preview_changeset(req: DocmatePreviewChangesetRequest) -> ToolResult:
    """Return preview cards for an existing changeset."""
    task_id = new_task_id("docmate")

    cs = _changeset_store.get(req.changeset_id)
    if not cs:
        return ToolResult(
            ok=False, tool="docmate_preview_changeset", task_id=task_id,
            summary=f"变更集不存在: {req.changeset_id}",
            error="changeset_not_found",
        )

    return ToolResult(
        ok=True, tool="docmate_preview_changeset", task_id=task_id,
        summary=f"变更预览：共 {len(cs['preview_cards'])} 项。",
        data={
            "changeset_id": req.changeset_id,
            "instruction": cs["instruction"],
            "preview_cards": cs["preview_cards"],
            "total_changes": len(cs["preview_cards"]),
        },
    )


# ---------------------------------------------------------------------------
# Apply changes
# ---------------------------------------------------------------------------

def docmate_apply_changeset(req: DocmateApplyChangesetRequest) -> ToolResult:
    """Apply accepted changes from a changeset to a new copy of the document."""
    task_id = new_task_id("docmate")

    cs = _changeset_store.get(req.changeset_id)
    if not cs:
        return ToolResult(
            ok=False, tool="docmate_apply_changeset", task_id=task_id,
            summary=f"变更集不存在: {req.changeset_id}",
            error="changeset_not_found",
        )

    doc_id = cs["doc_id"]
    doc_entry = _doc_store.get(doc_id)
    if not doc_entry:
        return ToolResult(
            ok=False, tool="docmate_apply_changeset", task_id=task_id,
            summary=f"文档状态已过期，请重新加载文档。",
            error="doc_not_found",
        )

    source_path = Path(doc_entry["source_path"])
    output_path = Path(req.save_as) if req.save_as else source_path.parent / f"{source_path.stem}_modified.docx"

    accepted_ids = set(req.accepted_change_ids)
    accepted_changes = [c for c in cs["changes"] if c["change_id"] in accepted_ids]
    rejected_changes = [c for c in cs["changes"] if c["change_id"] not in accepted_ids]

    if not accepted_changes:
        return ToolResult(
            ok=False, tool="docmate_apply_changeset", task_id=task_id,
            summary="没有选中任何变更，请至少选择一项后再应用。",
            error="no_changes_selected",
        )

    # Create backup
    backup_path = output_path.parent / f"{output_path.stem}_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        shutil.copy2(str(source_path), str(backup_path))
    except Exception as exc:
        return ToolResult(
            ok=False, tool="docmate_apply_changeset", task_id=task_id,
            summary=f"无法创建备份: {exc}",
            error=str(exc),
        )

    # Open a fresh copy and apply changes
    try:
        from docx import Document
        doc = Document(str(source_path))
    except Exception as exc:
        return ToolResult(
            ok=False, tool="docmate_apply_changeset", task_id=task_id,
            summary=f"无法打开文档: {exc}",
            error=str(exc),
        )

    applied_count = 0
    apply_errors: list[str] = []

    for change in accepted_changes:
        ctype = change.get("type", "text_replace")
        target = change.get("target", "")
        replacement = change.get("replacement", "")

        try:
            if ctype == "text_append":
                # Append new paragraph at end
                doc.add_paragraph(replacement)
                applied_count += 1

            elif ctype in ("text_replace", "text_delete"):
                applied = False
                # Search and replace in paragraphs
                for para in doc.paragraphs:
                    if target and target in para.text:
                        for run in para.runs:
                            if target in run.text:
                                run.text = run.text.replace(target, replacement, 1)
                                applied = True
                                break
                        if applied:
                            break

                # Also search tables
                if not applied and target:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if target in para.text:
                                        for run in para.runs:
                                            if target in run.text:
                                                run.text = run.text.replace(target, replacement, 1)
                                                applied = True
                                                break
                                        if applied:
                                            break
                                    if applied:
                                        break
                                if applied:
                                    break
                            if applied:
                                break
                        if applied:
                            break

                if applied:
                    applied_count += 1
                else:
                    apply_errors.append(f"未在文档中找到「{_truncate(target, 30)}」，跳过此项变更。")

        except Exception as exc:
            apply_errors.append(f"应用变更「{_truncate(target, 30)}」时出错: {exc}")

    # Save output
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
    except Exception as exc:
        return ToolResult(
            ok=False, tool="docmate_apply_changeset", task_id=task_id,
            summary=f"保存文档失败: {exc}",
            error=str(exc),
        )

    record_task_event(task_id, {
        "tool": "docmate_apply_changeset",
        "changeset_id": req.changeset_id,
        "applied": applied_count,
        "rejected": len(rejected_changes),
        "errors": len(apply_errors),
        "status": "ok",
    })

    return ToolResult(
        ok=True, tool="docmate_apply_changeset", task_id=task_id,
        summary=f"已应用 {applied_count} 项修改，拒绝 {len(rejected_changes)} 项。",
        data={
            "applied": applied_count,
            "rejected": len(rejected_changes),
            "errors": apply_errors,
            "output_path": str(output_path),
            "backup_path": str(backup_path),
            "source_path": str(source_path),
        },
    )


# ---------------------------------------------------------------------------
# Public interface expected by app.py / registry.py
#
# The application layer (app.py HTTP endpoints and registry.py tool listing)
# imports a stable interface: request models named *Request, async tool_*
# wrappers, and ALL_DOCMATE_SPECS. The pattern-matching implementation above
# is the source of truth; the wrappers below adapt it to that contract.
# ---------------------------------------------------------------------------

class ReadDocxRequest(BaseModel):
    file_path: str = Field(..., description=".docx 文件路径")


class GenerateChangesetRequest(BaseModel):
    doc_id: str = Field(..., description="文档 ID（来自 docmate_read_docx）")
    instruction: str = Field(..., description="用户自然语言编辑指令")
    context: Optional[Any] = Field(default=None, description="额外上下文（selected_text, form_type 等）")


class PreviewChangesetRequest(BaseModel):
    changeset_id: str = Field(..., description="ChangeSet ID")


class ApplyChangesetRequest(BaseModel):
    changeset_id: str = Field(..., description="ChangeSet ID")
    accepted_change_ids: list[str] = Field(default_factory=list, description="接受的变更 ID 列表")
    save_as: str = Field(default="", description="输出文件路径；留空则在源文件旁生成 *_modified.docx")


def _context_to_str(context: Any) -> str:
    if context is None:
        return ""
    if isinstance(context, str):
        return context
    try:
        return json.dumps(context, ensure_ascii=False)
    except (TypeError, ValueError):
        return str(context)


async def tool_docmate_read_docx(request: ReadDocxRequest) -> ToolResult:
    return docmate_read_docx(DocmateReadDocxRequest(file_path=request.file_path))


async def tool_docmate_generate_changeset(request: GenerateChangesetRequest) -> ToolResult:
    return docmate_generate_changeset(
        DocmateGenerateChangesetRequest(
            doc_id=request.doc_id,
            instruction=request.instruction,
            context=_context_to_str(request.context),
        )
    )


async def tool_docmate_preview_changeset(request: PreviewChangesetRequest) -> ToolResult:
    return docmate_preview_changeset(
        DocmatePreviewChangesetRequest(changeset_id=request.changeset_id)
    )


async def tool_docmate_apply_changeset(request: ApplyChangesetRequest) -> ToolResult:
    return docmate_apply_changeset(
        DocmateApplyChangesetRequest(
            changeset_id=request.changeset_id,
            accepted_change_ids=request.accepted_change_ids,
            save_as=request.save_as,
        )
    )


ALL_DOCMATE_SPECS = [
    ToolSpec(
        name="docmate_read_docx",
        description="解析 .docx 文件结构（段落、表格、图片统计），返回 doc_id 供后续编辑使用。",
        input_schema={
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Absolute path to .docx file"},
            },
            "required": ["file_path"],
        },
    ),
    ToolSpec(
        name="docmate_generate_changeset",
        description="根据自然语言指令分析已加载文档并生成修改方案（changeset）。",
        input_schema={
            "type": "object",
            "properties": {
                "doc_id": {"type": "string"},
                "instruction": {"type": "string"},
                "context": {"type": ["object", "string", "null"]},
            },
            "required": ["doc_id", "instruction"],
        },
    ),
    ToolSpec(
        name="docmate_preview_changeset",
        description="返回一个 changeset 的变更预览卡片。",
        input_schema={
            "type": "object",
            "properties": {
                "changeset_id": {"type": "string"},
            },
            "required": ["changeset_id"],
        },
    ),
    ToolSpec(
        name="docmate_apply_changeset",
        description="将选中的变更应用到文档副本（自动备份原文件）。",
        input_schema={
            "type": "object",
            "properties": {
                "changeset_id": {"type": "string"},
                "accepted_change_ids": {"type": "array", "items": {"type": "string"}},
                "save_as": {"type": "string"},
            },
            "required": ["changeset_id"],
        },
    ),
]
