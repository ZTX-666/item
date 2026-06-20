# -*- coding: utf-8 -*-
"""Generate 数字员工社区建设方案 Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date


def set_cell_shading(cell, fill):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return h


def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D9E2F3")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ===== Cover =====
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("数字员工社区建设方案\n（飞书版）")
    r.bold = True
    r.font.size = Pt(26)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("建筑企业 AI 智能体共创社区 · 执行手册")
    sr.font.size = Pt(14)
    sr.font.name = "微软雅黑"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(f"版本：V1.0    日期：{date.today().strftime('%Y年%m月%d日')}\n密级：内部资料")
    ir.font.size = Pt(11)
    ir.font.name = "宋体"
    ir._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    # ===== TOC placeholder =====
    add_heading(doc, "目  录", 1)
    toc_items = [
        "一、方案背景与建设目标",
        "二、国内外标杆实践参考",
        "三、社区定位与总体架构（飞书版）",
        "四、组织体系与角色职责",
        "五、飞书平台搭建指南",
        "六、社区运营机制（悬赏·开源·共创）",
        "七、社区行为准则（三大纪律八项注意）",
        "八、数据安全与合规治理",
        "九、建筑企业优先场景清单",
        "十、分阶段实施路线图",
        "十一、考核指标与激励机制",
        "十二、管理范式变革配套建议",
        "附录A：飞书资源清单模板",
        "附录B：悬赏任务书模板",
        "附录C：智能体上架审核表",
        "附录D：Champion 招募与职责说明",
        "附录E：90天启动 Checklist",
    ]
    for item in toc_items:
        add_para(doc, item)
    doc.add_page_break()

    # ===== Section 1 =====
    add_heading(doc, "一、方案背景与建设目标", 1)

    add_heading(doc, "1.1 背景", 2)
    add_para(doc, '领导提出：鼓励成员分享智能体，采用开源协作或悬赏机制共同完善；制定简单规则，先让系统运转起来，由本次会议参与者作为"元老"在线下推动建设。同时指出：AI 对生产力的提升必将倒逼管理范式变革；年轻人 AI 能力快速进化、个人生产力显著提升，但现有岗位体系（生产关系）尚无法有效消化这一跃迁。', indent=True)

    add_heading(doc, "1.2 核心判断", 2)
    add_bullet(doc, '社区不是"聊天群"，而是企业"数字员工资产库"的生产与流通系统。')
    add_bullet(doc, '建筑行业 AI 应用整体仍处于起步阶段，成功路径是"边界清晰的小场景 + 可复用资产 + 人机协同"，而非一步到位的大平台。')
    add_bullet(doc, '公司已使用飞书，应充分复用飞书 Aily 智能体平台、知识库、多维表格、群聊机器人等能力，降低建设与运营成本。')
    add_bullet(doc, '社区运转的关键不是技术炫度，而是"真实需求 → 任务交付 → 资产沉淀 → 跨部门复用"的闭环。')

    add_heading(doc, "1.3 建设目标", 2)
    add_table(doc, ["阶段", "时间", "目标"], [
        ["启动期", "0-3个月", "社区运转起来：3个标杆智能体、20名 Champion、500+ 成员入群"],
        ["沉淀期", "4-6个月", "形成资产库：30+ 可复用 Prompt/Workflow、5条业务线覆盖"],
        ["规模期", "7-12个月", "业务价值可量化：节约工时可统计、2-3个场景跨项目复制"],
        ["成熟期", "12个月+", "管理范式试点：人机协同流程、新角色通道、内部孵化机制"],
    ], [3, 3, 10])

    add_heading(doc, "1.4 建设原则", 2)
    add_bullet(doc, '先跑起来，再完善：规则从简，迭代优化。')
    add_bullet(doc, '业务驱动，价值可见：每个智能体必须对应真实痛点。')
    add_bullet(doc, '共建共享，贡献可见：开源协作 + 悬赏激励 + 贡献者榜。')
    add_bullet(doc, '安全合规，人机协同：敏感数据不上公网，关键节点人工确认。')
    add_bullet(doc, '飞书原生：在员工日常工作的 IM 环境中完成学习、搭建、使用、分享。')

    # ===== Section 2 =====
    add_heading(doc, "二、国内外标杆实践参考", 1)

    add_heading(doc, "2.1 国际建筑/工程企业", 2)

    add_heading(doc, "2.1.1 Balfour Beatty（英国，约1.3万人）", 3)
    add_para(doc, '投入约720万英镑部署 Microsoft Copilot，12周内完成46场培训、6345人参与，使用率超50%。核心做法：', indent=True)
    add_bullet(doc, '建立 50 名 Copilot Leader + 490 名 Champion 的分层网络')
    add_bullet(doc, '高管带头学习：每位高管与一名学徒一对一学习')
    add_bullet(doc, '多通道传播：Copilot Hub、FAQ、辟谣专栏、CEO 博客')
    add_bullet(doc, '可借鉴：分层 Champion 网络 + 高管示范 + 可量化采纳率')

    add_heading(doc, "2.1.2 Sisk（爱尔兰大型承包商）", 3)
    add_para(doc, '通过 4-35 人的用户组，3-4 周一轮实验，跨职能分享；成熟后按职能拆成 4-5 人深潜组。', indent=True)
    add_bullet(doc, 'Peer learning 优于自上而下命令')
    add_bullet(doc, '用教育管理 Shadow AI，而非单纯封堵')
    add_bullet(doc, '可借鉴：短周期实验 + 跨职能用户组')

    add_heading(doc, "2.1.3 Eiffage（法国，约7.8万人）", 3)
    add_para(doc, '与 Google Cloud 合作，限定 15 个优先场景；80 分钟必修 e-learning 才能开通内部 GenAI；员工 Prompt 审核后全集团共享。', indent=True)
    add_bullet(doc, '培训换权限：CEO/CFO 拍板，未完成培训无访问权')
    add_bullet(doc, '场景克制：按 adoption 选场景，不按技术炫度')
    add_bullet(doc, '可借鉴：Prompt 库全集团共享 + 必修培训 + 15 场景上限')

    add_heading(doc, "2.2 国内建筑企业", 2)

    add_heading(doc, "2.2.1 北京城建 / 城建智控", 3)
    add_bullet(doc, '建设"智能体助手广场"，覆盖工地监控、物料调度等场景')
    add_bullet(doc, '难点：现场数字化基础弱、专业场景（支模验算等）通用 AI 难理解')
    add_bullet(doc, '启示：先做管理类/办公类场景，现场强专业场景放第二期')

    add_heading(doc, "2.2.2 上海建工四建集团", 3)
    add_bullet(doc, '"云工大模型"（建筑业 MaaS）+ "四维·集智"多智能体集群')
    add_bullet(doc, '运维场景：能耗降低 5%、工单效率提升 20%')
    add_bullet(doc, '启示：从单点工具走向多智能体 + 数字孪生 + API 执行')

    add_heading(doc, "2.2.3 中国建筑国际等", 3)
    add_bullet(doc, '私有化大模型 + 岗位级 AI 助理（审单、造价、安全等）')
    add_bullet(doc, '启示：智能体按岗位/场景命名，便于员工理解和使用')

    add_heading(doc, "2.3 互联网/科技大厂社区模式", 2)

    add_table(doc, ["企业/平台", "模式", "核心机制", "可借鉴点"], [
        ["GitLab", "Hub-Spoke-Champion", "平台 Hub 管治理；各业务线 ATO；5-10% Champion", "三层架构清晰，职责不混"],
        ["GitHub", "AI Advocates + CoP", "志愿者 Champion + 社区自治 + 知识不锁私聊", "从中心团队过渡到社区自治"],
        ["腾讯元器/ADP", "智能体平台 + 企业级治理", "空间级权限、审计日志、评测工具", "发布前评测、发布后监控"],
        ["字节 Coze", "低代码 + 工作流 + 多渠道", "Workflow 编排、RAG 知识库、飞书一键发布", "降低搭建门槛，快速验证"],
        ["阿里云社区方法论", "任务池 + 案例库", "需求→任务→交付→复用闭环", "社区=项目制实践系统"],
        ["飞书 Aily", "企业级 Agent 平台", "自定义智能体、Workflow、MCP、群聊进化", "与飞书深度集成，天然适配"],
        ["Core Foundry", "内部开源", "PR 提交 Skill、Issue 悬赏、贡献者名录", "开源精神 + 内部闭环"],
        ["华为 FAB", "智能体加速器", "场景工作流预集成、开发周期从月到周", "工程化能力沉淀为模板"],
        ["神州信息 AI 创新中心", "四角色共创", "业务专家+产品经理+架构师+工程师同桌", "智能体是「养」出来的，不是一次性交付"],
    ], [2.5, 2.5, 4.5, 5])

    add_heading(doc, "2.4 标杆共性总结", 2)
    add_bullet(doc, '组织：Hub（平台）+ 业务 Champion（推广）+ 元老/Steering（方向）')
    add_bullet(doc, '机制：培训换权限、贡献可兑换、Show & Tell 常态化')
    add_bullet(doc, '资产：Prompt 库 + Workflow 智能体 + 领域知识包')
    add_bullet(doc, '度量：复用次数、节省工时、采纳率，而非登录次数')
    add_bullet(doc, '安全：内部工具替代 Shadow AI、分级管控、Human-in-the-loop')

    add_heading(doc, "2.5 更多大企业智能体社群参考", 2)

    add_table(doc, ["企业", "社群/平台模式", "关键做法", "来源/启示"], [
        ["交通银行 + 华为", "联合创新展示中心 + FAB 加速器", "战略协同、场景探索、技术创新、人才发展四维联动；智能体开发周期从月缩短到周", "金融同业交流标杆；建筑企业可类比「项目部+数字化」联合共创"],
        ["神州信息 AI 创新中心", "四角色同桌共创", "业务专家、AI 产品经理、智能体架构师、AI 工程师同室办公；智能体一版版试、一句句改", "智能体是「养」出来的；每个智能体近 50 个覆盖前中后台"],
        ["华为金融 FAB", "智能体加速器 + 12 大场景方案", "预集成工作流与工具链；AI 人才发展计划 3 年培养万名复合人才", "工程化模板 + 人才梯队同步建设"],
        ["腾讯 ADP / 元器", "空间级权限 + 审计 + 评测", "ADP 面向复杂业务；元器面向轻量探索与生态分发", "社区上架智能体需配套评测与审计机制"],
        ["字节 Coze / 扣子", "Workflow + 知识库 + 飞书发布", "零代码编排；Trace 调试；任务从对话工具升级为交付系统", "Champion 层级的首选搭建工具（可与 Aily 互补）"],
        ["阿里云开发者社区", "任务池 + 案例库 + 模板库", "需求→任务→交付→复用；五类基础智能体（需求分析、资料整理、内容生成、客服、复盘）", "社区运营=SOP 化项目制，非内容分享"],
        ["Core Foundry（开源范式）", "内部 Git + PR + Issue 悬赏", "Skills/Workflows/MCP 军械库；贡献勋章", "开源协作精神移植到企业内部"],
        ["GitHub 企业 playbook", "AI Advocates + CoP 双轮", "Advocates 催化 → CoP 化学反应；GitHub Discussions 沉淀知识", "Phase 3 社区自治，降低 Hub 依赖"],
        ["GitLab Hub-Spoke-Hub", "平台 Hub + ATO + Champion", "5-10% 员工作 Champion；用业务指标而非技术指标衡量成功", "组织模型可直接映射到飞书架构"],
    ], [2.5, 3, 4.5, 5.5])

    add_heading(doc, "2.6 建筑企业选型建议（飞书用户）", 2)
    add_bullet(doc, '基础层：全员开通「Aily 工作助手」+ 60-90 分钟必修培训（参考 Eiffage 培训换权限）')
    add_bullet(doc, '共建层：社区核心用「Aily 自定义智能体」+ 飞书知识库 + 多维表格（任务池/目录/悬赏）')
    add_bullet(doc, '进阶层：IT 团队用 Workflow + MCP/自定义连接器对接项目管理系统')
    add_bullet(doc, '互补层：个人探索可用 Coze 原型验证，上线版统一迁入 Aily 治理')

    # ===== Section 3 =====
    add_heading(doc, "三、社区定位与总体架构（飞书版）", 1)

    add_heading(doc, "3.1 社区定位", 2)
    add_para(doc, '全称建议：「数字员工共创社区」（Digital Worker Co-Creation Community）。定位：公司级 AI 智能体共创、共享、悬赏、运营平台，依托飞书实现"在 IM 里完成从学习到使用的全流程"。', indent=True)

    add_heading(doc, "3.2 三类核心资产", 2)
    add_table(doc, ["资产类型", "说明", "飞书载体", "示例"], [
        ["Prompt / Skill 包", "可复制的提示词与操作步骤", "飞书文档 + 多维表格目录", "安全巡检报告生成、招标文件摘要"],
        ["Workflow 智能体", "可重复运行的自动化流程", "飞书 Aily 自定义智能体", "会议纪要→待办→任务自动创建"],
        ["领域知识包", "规范、模板、案例（脱敏）", "飞书知识库 / 知识空间", "公司公文模板、安全规范 FAQ"],
    ], [3, 4, 3.5, 5])

    add_heading(doc, "3.3 飞书技术架构", 2)
    add_para(doc, '推荐"飞书原生四层架构"：', indent=True)
    add_bullet(doc, '交互层：飞书群聊（主社区群 + 条线分群）、飞书机器人、Aily 智能体单聊/群聊')
    add_bullet(doc, '应用层：飞书 Aily（自定义智能体 + Workflow）、Aily 工作助手（个人效率）、妙搭 OpenClaw（进阶定制）')
    add_bullet(doc, '资产层：飞书知识库（社区/wiki）、多维表格（任务池/贡献榜/悬赏池/智能体目录）')
    add_bullet(doc, '治理层：飞书管理后台（权限/审计）、社区规则文档、上架审核流程')

    add_heading(doc, "3.4 架构示意", 2)
    add_para(doc, 'Steering（元老+分管领导）→ 社区 Hub（2-3人兼职）→ 条线 Champion（每部门1-2人）→ 全员成员。Hub 统一管理：飞书知识库、Aily 应用中心、多维表格任务池、悬赏与积分。Champion 负责条线推广、需求收集、智能体试用反馈。成员贡献 Prompt/智能体 → 审核 → 上架 → 复用 → 积分/悬赏。', indent=True)

    add_heading(doc, "3.5 与领导原方案的对应", 2)
    add_table(doc, ["领导建议", "本方案落地"], [
        ["鼓励分享智能体", "飞书 Aily 应用中心 + 智能体目录（多维表格）+ Show & Tell"],
        ["开源协作", "飞书文档共建 + 版本记录 + 贡献者署名"],
        ["悬赏机制", "飞书多维表格悬赏池 + 审批流 + 积分兑换"],
        ["简单规则", "三大纪律八项注意（第七章）+ 1页纸速查"],
        ["元老推动", "Steering 6-12个月启动使命，之后转 Mentor"],
        ["管理范式变革", "第十二章配套：新角色、项目制、人机协同流程"],
        ["消化生产力跃迁", "贡献→积分→新角色/项目制/内部孵化通道"],
    ], [4, 12])

    # ===== Section 4 =====
    add_heading(doc, "四、组织体系与角色职责", 1)

    add_heading(doc, "4.1 组织架构", 2)
    add_table(doc, ["层级", "组成", "人数建议", "职责"], [
        ["Steering 指导组", "分管领导 + 会议「元老」 + 数字化负责人", "5-7人", "定方向、批预算、解冲突、季度复盘"],
        ["社区 Hub 运营组", "数字化+人力+业务骨干（兼职）", "2-3人", "平台运维、悬赏管理、活动组织、度量统计"],
        ["Champion 网络", "各部门自愿+推荐", "员工总数1-2%", "推广、带教、需求收集、试用反馈"],
        ["全体会员", "完成入门培训的员工", "逐步扩大", "使用、反馈、贡献、复用"],
    ], [2.5, 3.5, 2.5, 7])

    add_heading(doc, "4.2 角色详细说明", 2)

    add_heading(doc, "4.2.1 元老（Steering 成员）", 3)
    add_bullet(doc, '使命周期：6-12 个月，之后自动转 Mentor 或退出，避免"终身特权"')
    add_bullet(doc, '线下职责：部门宣贯、资源协调、跨部门破冰、参与季度评审')
    add_bullet(doc, '线上职责：飞书群置顶规则、重大悬赏终审、Show & Tell 嘉宾')
    add_bullet(doc, '时间投入：约 2-4 小时/月')

    add_heading(doc, "4.2.2 社区 Hub 运营", 3)
    add_bullet(doc, '维护飞书知识库结构、多维表格、Aily 应用中心')
    add_bullet(doc, '组织双周 Show & Tell、月度贡献者榜')
    add_bullet(doc, '审核智能体上架、管理悬赏流程')
    add_bullet(doc, '对接 IT/信息安全，处理合规问题')
    add_bullet(doc, '时间投入：约 4-8 小时/周（可轮值）')

    add_heading(doc, "4.2.3 Champion（智能体推广大使）", 3)
    add_bullet(doc, '本部门 AI 第一联系人：同事有问题找 Champion，而非直接找 IT')
    add_bullet(doc, '收集本部门需求，提交到飞书多维表格「需求池」')
    add_bullet(doc, '试点新智能体，收集反馈，协助 Hub 调优')
    add_bullet(doc, '组织本部门 mini 分享（30分钟）')
    add_bullet(doc, '时间投入：约 30-60 分钟/周')
    add_bullet(doc, '详见附录 D')

    add_heading(doc, "4.3 Champion 覆盖要求", 2)
    add_table(doc, ["业务条线", "Champion 配置", "优先场景"], [
        ["工程/项目部", "每大型项目/区域 1 人", "安全报告、进度周报、技术方案初稿"],
        ["商务/成本", "每部门 1-2 人", "招标摘要、合同条款比对、变更测算"],
        ["设计/技术", "每部门 1 人", "规范检索、图纸说明、技术交底"],
        ["安全/质量", "每部门 1 人", "巡检记录、隐患分析、整改通知"],
        ["职能（人力/财务/行政）", "每部门 1 人", "制度问答、报销指引、会议纪要"],
        ["数字化/IT", "2-3 人", "Aily 搭建支持、连接器开发、平台运维"],
    ], [3, 4, 8])

    # ===== Section 5 =====
    add_heading(doc, "五、飞书平台搭建指南", 1)

    add_heading(doc, "5.1 飞书资源清单（Day 1 必建）", 2)
    add_table(doc, ["资源", "名称建议", "用途"], [
        ["飞书群", "「数字员工共创社区」主群", "公告、Show & Tell、日常交流"],
        ["飞书群", "「数字员工·工程/商务/职能」分群", "条线深潜讨论（Phase 2 再建）"],
        ["知识库", "数字员工社区 Wiki", "规则、教程、案例、FAQ"],
        ["多维表格", "智能体资产目录", "名称、场景、贡献者、复用次数、链接"],
        ["多维表格", "需求/任务池", "业务需求收集、任务拆解、状态跟踪"],
        ["多维表格", "悬赏池", "赏金任务发布、认领、交付、验收"],
        ["多维表格", "贡献者积分榜", "积分规则、排名、兑换记录"],
        ["Aily", "社区官方助手（Hub 维护）", "新人引导、规则问答、目录检索"],
        ["Aily", "3 个标杆自定义智能体", "第一期交付物（见第九章）"],
        ["文档", "1 页纸《社区速查手册》", "入群必读，置顶链接"],
    ], [2.5, 4.5, 8.5])

    add_heading(doc, "5.2 飞书 Aily 智能体搭建路径", 2)
    add_para(doc, '根据飞书官方能力分层，建议社区采用以下选型：', indent=True)
    add_table(doc, ["层级", "飞书产品", "适用场景", "社区角色"], [
        ["个人效率", "Aily 工作助手", "写文档、查资料、做总结", "全员默认开通，入门体验"],
        ["团队共建", "Aily 自定义智能体", "团队知识问答、Workflow、定时任务", "社区核心交付形态"],
        ["进阶定制", "妙搭 OpenClaw / MCP", "接入内部系统、复杂插件", "IT + 高阶 Champion 使用"],
    ], [2, 3, 4.5, 5.5])

    add_heading(doc, "5.3 自定义智能体标准搭建六步（SOP）", 2)
    add_bullet(doc, 'Step 1 场景定义：明确输入、输出、使用人、频率、成功标准（用附录 C 填写）')
    add_bullet(doc, 'Step 2 知识接入：关联飞书知识空间/文档，确保知识可更新')
    add_bullet(doc, 'Step 3 人设与 Prompt：用 Aily Prompt IDE 编写，含角色、技能、限制')
    add_bullet(doc, 'Step 4 Workflow 编排（如需）：LLM → 条件分支 → 飞书多维表格写入 → 消息通知')
    add_bullet(doc, 'Step 5 调试与用例：至少 10 条测试用例，保存优质对话为 Example')
    add_bullet(doc, 'Step 6 发布与进化：发布到群聊，开启"与团队一起进化"，记录版本变更')

    add_heading(doc, "5.4 飞书群聊运营配置", 2)
    add_bullet(doc, '群公告：社区定位 + 三大纪律链接 + 入群培训链接')
    add_bullet(doc, '群机器人：Aily 社区助手 + 悬赏提醒机器人（多维表格自动化）')
    add_bullet(doc, '群置顶：①速查手册 ②智能体目录 ③悬赏池 ④Show & Tell 日历')
    add_bullet(doc, '群标签/分组：#需求 #成果 #求助 #悬赏 #ShowTell')
    add_bullet(doc, '自动化：新成员入群 → 自动发送欢迎卡片 + 培训链接 + 3个推荐智能体')

    add_heading(doc, "5.5 多维表格关键字段设计", 2)

    add_heading(doc, "5.5.1 智能体资产目录", 3)
    add_bullet(doc, '字段：名称 | 场景分类 | 贡献者 | 部门 | Aily 链接 | 知识库链接 | 复用次数 | 状态（草稿/审核/上架/下线）| 审核人 | 更新时间')

    add_heading(doc, "5.5.2 需求/任务池", 3)
    add_bullet(doc, '字段：需求描述 | 提出人 | 部门 | 优先级 | 拆解任务 | 认领人 | 状态 | 关联智能体 | 完成日期')

    add_heading(doc, "5.5.3 悬赏池", 3)
    add_bullet(doc, '字段：赏金标题 | 场景 | 赏金积分 | 发布人 | 认领人 | 交付物链接 | 验收标准 | 状态 | 截止日期')

    add_heading(doc, "5.6 与现有系统集成（Phase 2+）", 2)
    add_bullet(doc, '飞书 Aily 自定义连接器：封装 OA、项目管理系统 HTTP API')
    add_bullet(doc, '飞书项目 MCP：覆盖工作项管理、待办、度量分析（40+ 工具）')
    add_bullet(doc, '飞书开放平台：群聊总结、消息卡片、审批流联动')

    # ===== Section 6 =====
    add_heading(doc, "六、社区运营机制（悬赏·开源·共创）", 1)

    add_heading(doc, "6.1 运营闭环：需求→任务→交付→复用", 2)
    add_para(doc, '参考阿里云社区方法论，社区本质是"项目制实践系统"，而非内容分享群。闭环如下：', indent=True)
    add_bullet(doc, '需求收集：Champion 收集 → 写入飞书「需求池」→ Hub 每周汇总')
    add_bullet(doc, '任务拆解：Hub 或"需求分析"智能体将模糊需求拆为标准任务')
    add_bullet(doc, '任务认领：悬赏发布 或 自愿认领 → 指定交付物和截止日期')
    add_bullet(doc, '交付验收：Hub + 业务方验收 → 上架资产目录 → 发放积分')
    add_bullet(doc, '复用推广：Show & Tell 展示 → 写入知识库 → 统计复用次数')

    add_heading(doc, "6.2 悬赏机制设计", 2)
    add_table(doc, ["类型", "示例", "积分范围", "验收标准"], [
        ["需求悬赏", "安全周报从2小时压到30分钟", "500-2000", "3人试用达标 + 文档完整"],
        ["资产悬赏", "提交可复用 Workflow 智能体", "1000-5000", "通过审核上架 + 5次复用"],
        ["攻坚悬赏", "打通 OA 审批 + 飞书通知", "3000-10000", "端到端跑通 + IT 安全确认"],
        ["教学悬赏", "制作 15 分钟入门微课", "300-800", "播放 50 次 + 评分 4+/5"],
        ["Bug 悬赏", "修复某智能体幻觉/格式问题", "200-1000", "10 条测试用例通过"],
    ], [2, 4.5, 2, 6])

    add_heading(doc, "6.3 积分与兑换", 2)
    add_table(doc, ["行为", "积分"], [
        ["完成必修培训", "+50"],
        ["首次贡献 Prompt（审核通过）", "+100"],
        ["首次上架智能体", "+500"],
        ["智能体被复用 10 次", "+200"],
        ["完成悬赏任务", "按赏金表"],
        ["Show & Tell 分享", "+150"],
        ["帮助同事（被 @感谢并确认）", "+30"],
    ], [8, 4])
    add_para(doc, '兑换建议：培训名额、书籍、购物卡（100-500元档）、评优加分、"数字员工产品经理"试点项目资格。', indent=True)

    add_heading(doc, "6.4 开源协作机制", 2)
    add_bullet(doc, '所有 Prompt/Workflow 存飞书文档，开启"允许评论"和"版本历史"')
    add_bullet(doc, '改进流程：Fork（复制）→ 修改 → 提交 Hub 审核 → 合并入主版本 → 贡献者署名')
    add_bullet(doc, 'Issue 机制：在「需求池」提 Issue，标注 #待解决，Champion 可认领')
    add_bullet(doc, '贡献者榜：每月飞书群 + 知识库发布 Top 10')

    add_heading(doc, "6.5 活动日历（标准 SOP）", 2)
    add_table(doc, ["频率", "活动", "时长", "负责人", "飞书形式"], [
        ["入群时", "必修培训 + 推荐 3 个智能体", "60-90 min", "Hub + 人力", "飞书学堂/文档 + 测验"],
        ["双周", "Show & Tell", "30 min", "轮值 Champion", "飞书视频会议 + 录屏存知识库"],
        ["每月", "贡献者榜 + 悬赏公示", "—", "Hub", "群公告 + 多维表格"],
        ["每季", "Steering 复盘", "2 h", "Steering", "线下 + 飞书纪要"],
        ["按需", "条线 Mini Hackathon", "半天", "Champion", "线下 + 飞书群直播"],
    ], [1.5, 3.5, 1.5, 2.5, 6.5])

    add_heading(doc, "6.6 新人 onboarding 流程", 2)
    add_bullet(doc, 'Step 1：扫码/邀请入群 → 自动欢迎消息')
    add_bullet(doc, 'Step 2：完成 60-90 分钟必修培训（Prompt 基础 + 数据安全）')
    add_bullet(doc, 'Step 3：试用 3 个推荐智能体（安全报告/会议纪要/制度问答）')
    add_bullet(doc, 'Step 4：在群里做一次"我的第一次 AI 体验"分享（可选，+30 积分）')
    add_bullet(doc, 'Step 5：开通 Aily 工作助手完整权限')

    # ===== Section 7 =====
    add_heading(doc, "七、社区行为准则（三大纪律八项注意）", 1)

    add_heading(doc, "三大纪律", 2)
    add_para(doc, '1. 不上传未脱敏的项目数据、合同、图纸至公网 AI（ChatGPT 等）；公司内部只用飞书 Aily 等合规工具。', bold=True)
    add_para(doc, '2. 涉及安全责任、签章、结算、法律效力的内容，必须人工审核，智能体输出仅供参考。', bold=True)
    add_para(doc, '3. 有好用的 Prompt、智能体、工作流，应共享到社区资产库，不得长期私藏。', bold=True)

    add_heading(doc, "八项注意", 2)
    items = [
        "先小场景试点，再追求全自动——复杂现场场景需人机协同。",
        "贡献时写清楚：适用场景、操作步骤、输入示例、限制条件。",
        "复用他人资产请反馈效果（好评/问题），帮助持续改进。",
        "发现数据安全或合规风险，立即在群里 @Hub 运营组。",
        "不贬低尚未使用 AI 的同事；提倡一对一带教。",
        "悬赏以可验证交付为准，不议价、不拖延。",
        "工具会变，沉淀 SOP 和知识比绑定某个 App 更重要。",
        "元老是服务员和推广大使，不是审批官——审核标准公开透明。",
    ]
    for i, item in enumerate(items, 1):
        add_para(doc, f"{i}. {item}")

    add_heading(doc, "违规处理", 2)
    add_bullet(doc, '轻度（首次不当使用公网 AI）：警告 + 补训')
    add_bullet(doc, '中度（上传敏感数据）：暂停社区权限 + 通报部门')
    add_bullet(doc, '重度（造成安全/合规事件）：按公司信息安全制度处理')

    # ===== Section 8 =====
    add_heading(doc, "八、数据安全与合规治理", 1)

    add_heading(doc, "8.1 数据分级", 2)
    add_table(doc, ["级别", "范围", "AI 使用规则"], [
        ["L1 公开", "公司宣传、公开制度摘要", "可用任何合规工具"],
        ["L2 内部", "内部模板、脱敏案例、通用 SOP", "仅限飞书 Aily / 内部平台"],
        ["L3 机密", "合同、造价、未公开项目信息", "禁止输入任何 AI；人工处理"],
        ["L4 绝密", "战略、人事敏感、安全事故原始资料", "禁止输入；专项审批"],
    ], [2, 5, 8.5])

    add_heading(doc, "8.2 Human-in-the-loop 清单（必须人工确认）", 2)
    add_bullet(doc, '对外函件、律师函、索赔文件')
    add_bullet(doc, '安全整改通知书、事故报告')
    add_bullet(doc, '工程变更、签证、结算文件')
    add_bullet(doc, '人事决策、绩效评价')
    add_bullet(doc, '任何需加盖公章的输出')

    add_heading(doc, "8.3 Shadow AI 管理", 2)
    add_bullet(doc, '提供好用的内部替代（飞书 Aily），降低员工偷用公网工具动机')
    add_bullet(doc, '培训中明确：公网 AI 可能留存数据、产生幻觉、无法追责')
    add_bullet(doc, '参考 Eiffage：培训完成才开通权限')

    add_heading(doc, "8.4 智能体上架审核（五关）", 2)
    add_bullet(doc, '① 场景合规：是否涉及 L3/L4 数据')
    add_bullet(doc, '② 输出审核：是否有 Human-in-the-loop 节点')
    add_bullet(doc, '③ 测试用例：≥10 条，通过率 ≥80%')
    add_bullet(doc, '④ 权限范围：知悉范围、可使用人群')
    add_bullet(doc, '⑤ Hub + 业务方双签')

    # ===== Section 9 =====
    add_heading(doc, "九、建筑企业优先场景清单", 1)

    add_heading(doc, "9.1 第一期（0-3个月，建议 3 个标杆）", 2)
    add_table(doc, ["优先级", "场景", "智能体形态", "预期收益", "难度"], [
        ["P0", "项目安全/质量巡检报告生成", "Aily Workflow + 模板", "单次节约 1-2h", "低"],
        ["P0", "会议纪要 → 待办 → 飞书任务", "Aily 自定义智能体", "单次节约 30min", "低"],
        ["P0", "公司制度/流程 FAQ 问答", "Aily + 知识库", "减少重复咨询", "低"],
    ], [1.5, 4, 3.5, 3, 1.5])

    add_heading(doc, "9.2 第二期（4-6个月）", 2)
    add_table(doc, ["优先级", "场景", "说明"], [
        ["P1", "招标文件要点提取", "上传 PDF → 结构化摘要 + 合规 checklist"],
        ["P1", "施工方案/技术交底初稿", "基于模板 + 规范知识库生成"],
        ["P1", "工程周报/月报汇总", "多维表格数据 → 自动生成汇报"],
        ["P1", "变更签证测算辅助", "需 Human-in-the-loop，仅做初算"],
        ["P2", "合同条款比对", "两份合同差异高亮（法务审核）"],
    ], [1.5, 4, 10])

    add_heading(doc, "9.3 第三期（7-12个月，需系统集成）", 2)
    add_bullet(doc, '项目管理系统 + 飞书 Aily 连接器：进度预警、资源调度建议')
    add_bullet(doc, 'BIM/智慧工地数据 + 智能体：安全风险识别（需数字孪生基础）')
    add_bullet(doc, '多智能体协同：参考上海建工"四维·集智"模式（运维/后勤场景优先）')

    add_heading(doc, "9.4 场景选择原则", 2)
    add_bullet(doc, '高频 + 边界清晰 + 可量化 + 风险可控')
    add_bullet(doc, '优先办公/管理类，现场强专业放后')
    add_bullet(doc, '参考 Eiffage：同时推进不超过 15 个场景')

    # ===== Section 10 =====
    add_heading(doc, "十、分阶段实施路线图", 1)

    add_heading(doc, "10.1 Phase 0：准备期（第 1-2 周）", 2)
    checklist_p0 = [
        "成立 Steering（5-7人）并召开启动会",
        "任命 Hub 运营（2-3人兼职）",
        "创建飞书群、知识库、4 张多维表格",
        "发布《三大纪律八项注意》并置顶",
        "确定 3 个 P0 场景与负责人",
        "招募首批 15-20 名 Champion",
        "申请/确认飞书 Aily 专业版额度",
    ]
    for item in checklist_p0:
        add_bullet(doc, f"☐ {item}")

    add_heading(doc, "10.2 Phase 1：MVP 运转（第 3-12 周）", 2)
    checklist_p1 = [
        "上线 3 个标杆 Aily 智能体",
        "完成首批 Champion 培训（2h）",
        "举办 4 场 Show & Tell（双周一次）",
        "发布 3 个首期悬赏任务",
        "必修培训上线，目标 200 人完成",
        "智能体目录 ≥10 条（含 Prompt 包）",
        "Steering 第一次季度预览",
    ]
    for item in checklist_p1:
        add_bullet(doc, f"☐ {item}")

    add_heading(doc, "10.3 Phase 2：资产沉淀（第 4-6 月）", 2)
    add_bullet(doc, '每条业务线 ≥2 个上架智能体')
    add_bullet(doc, '建立条线分群（工程/商务/职能）')
    add_bullet(doc, '积分兑换机制正式运行')
    add_bullet(doc, '跨项目复制 1 个标杆场景')
    add_bullet(doc, '元老转 Mentor 机制启动')

    add_heading(doc, "10.4 Phase 3：机制升级（第 7-12 月）", 2)
    add_bullet(doc, '评估「智能体助手广场」形态（内部应用市场）')
    add_bullet(doc, '1 个管理范式试点（人机协同周例会）')
    add_bullet(doc, 'HR 联动：AI 能力纳入岗位发展参考')
    add_bullet(doc, '输出年度白皮书：场景、ROI、贡献者')
    add_bullet(doc, '启动 1 个内部孵化项目')

    # ===== Section 11 =====
    add_heading(doc, "十一、考核指标与激励机制", 1)

    add_heading(doc, "11.1 社区健康度指标", 2)
    add_table(doc, ["指标", "启动期目标", "6个月目标", "12个月目标"], [
        ["注册/入群人数", "500", "1500", "3000+"],
        ["完成必修培训比例", "30%", "60%", "80%"],
        ["活跃贡献者（月）", "20", "50", "100"],
        ["上架智能体数", "10", "30", "60+"],
        ["月均复用次数", "100", "500", "2000+"],
        ["Show & Tell 场次", "4", "12", "24"],
    ], [4, 3, 3, 3])

    add_heading(doc, "11.2 业务价值指标", 2)
    add_table(doc, ["指标", "度量方式"], [
        ["场景节约工时", "智能体目录 × 单次节约 × 月使用次数（抽样调研）"],
        ["需求交付率", "需求池关闭率 ≥60%"],
        ["Shadow AI 下降", "信息安全事件 / 问卷调研"],
        ["员工 AI 素养", "培训通过率 + 自评问卷"],
    ], [4, 12])

    add_heading(doc, "11.3 激励体系", 2)
    add_bullet(doc, '精神激励：月度/季度贡献者榜、Steering 颁奖、内部宣传')
    add_bullet(doc, '物质激励：积分兑换、悬赏奖金、评优加分')
    add_bullet(doc, '发展激励：优先参与数字化项目、新角色试点、外部培训名额')
    add_bullet(doc, '组织激励：部门上架智能体数量纳入数字化考核参考（非硬性 KPI）')

    # ===== Section 12 =====
    add_heading(doc, "十二、管理范式变革配套建议", 1)

    add_heading(doc, "12.1 回应「生产关系」矛盾的三条通道", 2)

    add_heading(doc, "通道一：新角色", 3)
    add_table(doc, ["角色", "职责", "来源"], [
        ["AI 训练师 / Prompt 工程师", "调优智能体、维护知识库", "Champion 进阶"],
        ["场景产品经理", "需求→智能体→ROI 闭环", "业务骨干 + 数字化"],
        ["人机协同流程设计师", " redesign 审批/报告/巡检流程", "部门负责人 + Champion"],
    ], [3.5, 5.5, 6.5])

    add_heading(doc, "通道二：项目制", 3)
    add_bullet(doc, '高产能年轻人可领"智能体 mini 项目"（如 2 周交付 1 个 Workflow）')
    add_bullet(doc, '考核按交付成果，而非仅按原岗位编制')
    add_bullet(doc, '参考神州信息"四角色同桌"：业务+产品+架构+工程')

    add_heading(doc, "通道三：内部创业/孵化", 3)
    add_bullet(doc, '社区验证成功的智能体 → 申请"孵化" → 配 1-2 人小团队深化')
    add_bullet(doc, '成功后可成为公司数字化产品或对外的智慧建造服务')
    add_bullet(doc, '参考上海建工"云工大模型"路径')

    add_heading(doc, "12.2 人机协同管理试点", 2)
    add_bullet(doc, '选 1 个项目部试点"人机协同周例会"：AI 生成进度/安全摘要 → 人决策')
    add_bullet(doc, '试点"AI 辅助汇报"：周报由智能体初稿 → 负责人 15 分钟修订')
    add_bullet(doc, '记录前后工时对比，作为管理范式变革的实证材料')

    add_heading(doc, "12.3 代际协作：年轻人 × 老师傅", 2)
    add_para(doc, '参考 Beard Construction：年轻员工数字化强但缺现场经验，现场管理者经验丰富但数字化弱。建议"双向 Mentor"配对，年轻人教 AI 工具，老师傅教业务判断，共同完善场景智能体。', indent=True)

    # ===== Appendices =====
    doc.add_page_break()
    add_heading(doc, "附录 A：飞书资源清单模板", 1)
    add_table(doc, ["序号", "资源类型", "名称", "链接/ID", "管理员", "状态"], [
        ["1", "飞书群", "数字员工共创社区", "（填写）", "Hub", "待建"],
        ["2", "知识库", "社区 Wiki", "（填写）", "Hub", "待建"],
        ["3", "多维表格", "智能体资产目录", "（填写）", "Hub", "待建"],
        ["4", "多维表格", "需求/任务池", "（填写）", "Hub", "待建"],
        ["5", "多维表格", "悬赏池", "（填写）", "Hub", "待建"],
        ["6", "多维表格", "贡献者积分榜", "（填写）", "Hub", "待建"],
        ["7", "Aily", "社区官方助手", "（填写）", "IT+Hub", "待建"],
        ["8", "文档", "社区速查手册", "（填写）", "Hub", "待建"],
    ], [1, 2.5, 3.5, 3, 2, 1.5])

    add_heading(doc, "附录 B：悬赏任务书模板", 1)
    add_para(doc, '【赏金任务书】', bold=True)
    template_fields = [
        "任务编号：BH-2026-___",
        "任务标题：________________________",
        "业务场景：□工程 □商务 □安全 □职能 □其他",
        "需求描述：（什么痛点、谁在用、现在怎么做）",
        "交付物：□ Prompt 包 □ Aily 智能体 □ Workflow □ 培训微课 □ 其他",
        "验收标准：（可量化，如：3人试用、10条测试用例、节约X小时）",
        "赏金积分：________",
        "发布人：________  发布日期：________",
        "认领人：________  截止日期：________",
        "交付链接：________",
        "验收人：________  验收日期：________  □通过 □退回",
        "验收意见：________________________",
    ]
    for f in template_fields:
        add_para(doc, f)

    add_heading(doc, "附录 C：智能体上架审核表", 1)
    add_table(doc, ["审核项", "要求", "审核人", "结果"], [
        ["场景合规", "不涉及 L3/L4 数据", "Hub", "□"],
        ["Human-in-the-loop", "关键输出需人工确认", "业务方", "□"],
        ["测试用例", "≥10条，通过率≥80%", "贡献者+Hub", "□"],
        ["权限范围", "知悉范围明确", "Hub", "□"],
        ["文档完整", "使用说明+限制条件", "Hub", "□"],
        ["双签批准", "Hub + 业务负责人", "Steering", "□"],
    ], [3, 5, 3, 2.5])

    add_heading(doc, "附录 D：Champion 招募与职责说明", 1)
    add_heading(doc, "招募条件", 2)
    add_bullet(doc, '对 AI 有兴趣，愿意分享，在部门内有影响力')
    add_bullet(doc, '已完成社区必修培训')
    add_bullet(doc, '部门负责人推荐 + 自愿报名')
    add_bullet(doc, '承诺每周 30-60 分钟社区投入')

    add_heading(doc, "职责清单", 2)
    add_bullet(doc, '担任本部门 AI 第一联系人')
    add_bullet(doc, '每月收集 ≥1 条业务需求到需求池')
    add_bullet(doc, '试点新智能体并提交反馈')
    add_bullet(doc, '每季组织 1 次部门 mini 分享（可联合其他 Champion）')
    add_bullet(doc, '参加 Hub 双月 Champion 短会（30 min）')

    add_heading(doc, "权益", 2)
    add_bullet(doc, 'Aily 专业版优先额度')
    add_bullet(doc, '外部培训/会议优先名额')
    add_bullet(doc, '贡献积分双倍（Champion 身份期间）')
    add_bullet(doc, '年度"数字员工之星"评选资格')

    add_heading(doc, "附录 E：90 天启动 Checklist", 1)
    weeks = [
        ("第 1 周", ["Steering 启动会", "飞书群+知识库+表格创建", "发布社区规则", "确定 Hub 运营"]),
        ("第 2 周", ["Champion 招募通知", "Aily 额度确认", "3个P0场景分工", "必修培训内容准备"]),
        ("第 3-4 周", ["Champion 培训（2h）", "启动 3 个智能体搭建", "首期悬赏发布", "入群推送"]),
        ("第 5-6 周", ["第1次 Show & Tell", "智能体内测", "收集反馈迭代", "培训推广"]),
        ("第 7-8 周", ["3个标杆智能体上架", "第2次 Show & Tell", "贡献者榜首发", "100人完成培训"]),
        ("第 9-10 周", ["第3次 Show & Tell", "条线需求汇总", "第2期悬赏", "跨部门复制1个场景"]),
        ("第 11-12 周", ["90天复盘会", "Steering 评审", "Phase 2 计划", "元老中期总结"]),
    ]
    for week, tasks in weeks:
        add_para(doc, week, bold=True)
        for t in tasks:
            add_bullet(doc, f"☐ {t}")

    # Footer note
    doc.add_paragraph()
    add_para(doc, '— 文档结束 —')
    add_para(doc, '本方案为内部执行参考，具体资源名称、积分额度、组织架构可根据公司实际情况调整。飞书 Aily 功能以官方最新版本为准。')

    return doc


if __name__ == "__main__":
    output_path = r"e:\3311 AI\数字员工社区建设方案（飞书版）.docx"
    document = build_document()
    document.save(output_path)
    print(f"Generated: {output_path}")
